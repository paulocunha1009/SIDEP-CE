# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(".")
OUT = ROOT / "outputs"
OUT.mkdir(exist_ok=True)

BLUE = RGBColor(0x0B, 0x35, 0x6D)
GREEN = RGBColor(0x00, 0x7A, 0x3D)
MUTED = RGBColor(0x55, 0x66, 0x7A)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(10)


def style_doc(doc, title):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, RGBColor(0x1F, 0x4D, 0x78)),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(6)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(title)
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def add_cover(doc, title, subtitle, version):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("SIDEP-CE")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(14)
    r.font.color.rgb = GREEN

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(24)
    r.font.color.rgb = BLUE

    p = doc.add_paragraph()
    r = p.add_run(subtitle)
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    r.font.color.rgb = MUTED

    table = doc.add_table(rows=3, cols=2)
    table.autofit = False
    for row in table.rows:
        row.cells[0].width = Inches(1.6)
        row.cells[1].width = Inches(4.8)
    rows = [
        ("Versão", version),
        ("Data", "08/07/2026"),
        ("Escopo", "Piloto controlado, banco de itens, avaliação diagnóstica, relatórios e preparação pré-TRI/TRI."),
    ]
    for i, (k, v) in enumerate(rows):
        set_cell_text(table.cell(i, 0), k, True)
        set_cell_text(table.cell(i, 1), v)
        set_cell_shading(table.cell(i, 0), "E8EEF5")

    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_manual():
    doc = Document()
    style_doc(doc, "Manual de Uso - SIDEP-CE v0.7")
    add_cover(
        doc,
        "Manual de Uso do SIDEP-CE",
        "Sistema de Diagnóstico da Educação Profissional do Ceará",
        "0.7",
    )

    doc.add_heading("1. Finalidade", level=1)
    doc.add_paragraph(
        "O SIDEP-CE é uma plataforma de avaliação diagnóstica da Educação Profissional do Ceará. "
        "O sistema organiza cursos técnicos, competências, descritores, banco de itens, avaliações, respostas e relatórios pedagógicos."
    )
    doc.add_paragraph(
        "O foco não é apenas aplicar prova online. O foco é transformar o desempenho dos estudantes em informação pedagógica para intervenção, recomposição e melhoria da aprendizagem técnica."
    )

    doc.add_heading("2. Perfis de Acesso", level=1)
    table = doc.add_table(rows=1, cols=2)
    set_cell_text(table.cell(0, 0), "Perfil", True)
    set_cell_text(table.cell(0, 1), "Acesso principal", True)
    set_cell_shading(table.cell(0, 0), "E8EEF5")
    set_cell_shading(table.cell(0, 1), "E8EEF5")
    rows = [
        ("Aluno", "Acessa somente a prova por código e nome completo."),
        ("Professor técnico", "Valida itens, cria avaliações e acompanha relatórios das suas aplicações."),
        ("Coordenador/professor técnico", "Acompanha curso, curadoria e uso pedagógico dos dados."),
        ("Gestão escolar", "Acessa professores, avaliações e relatórios da própria escola."),
        ("CREDE/SEFOR", "Acessa escolas e relatórios da própria regional."),
        ("SEDUC", "Acompanha indicadores agregados da rede."),
        ("Administrador master", "Possui acesso total e pode redefinir senhas e executar migrações."),
    ]
    for perfil, acesso in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], perfil, True)
        set_cell_text(cells[1], acesso)

    doc.add_heading("3. Banco de Itens", level=1)
    doc.add_paragraph(
        "O banco de itens é organizado por curso técnico, componente curricular, competência, descritor e questão. "
        "Competência é uma capacidade ampla; descritor é a evidência avaliável; questão é o item que mede o descritor."
    )
    add_bullets(
        doc,
        [
            "Questões em rascunho ou revisão não entram em avaliação.",
            "Somente questões validadas podem ser usadas em prova.",
            "O professor pode abrir a questão em modal para leitura completa antes de validar.",
            "O sistema bloqueia cadastro de questões duplicadas.",
        ],
    )

    doc.add_heading("4. Exportação Curricular", level=1)
    doc.add_paragraph(
        "Na área Banco de Itens, o usuário pode exportar cursos, componentes, competências e descritores em Markdown ou PDF. "
        "A exportação pode ser feita para todos os cursos ou para um curso específico."
    )

    doc.add_heading("5. Criação e Aplicação da Avaliação", level=1)
    add_bullets(
        doc,
        [
            "A avaliação deve ter no mínimo 20 e no máximo 80 questões.",
            "O código da avaliação é gerado automaticamente pelo sistema.",
            "Após a abertura, o código fica imutável.",
            "Código usado ou excluído não pode ser reaproveitado.",
            "A prova não pode conter questões duplicadas ou de contexto muito semelhante.",
            "A ordem das questões é embaralhada para cada estudante.",
        ],
    )

    doc.add_heading("6. Relatórios e Pré-TRI", level=1)
    doc.add_paragraph(
        "O sistema calcula acertos, percentual bruto, desempenho por descritor, componente e competência. "
        "Na fase atual, esses resultados são pré-TRI. A TRI plena dependerá de volume de respostas, itens âncora e calibração psicométrica."
    )

    doc.add_heading("7. Uso Online e Cuidados", level=1)
    add_bullets(
        doc,
        [
            "Configurar VITE_SUPABASE_URL e VITE_SUPABASE_PUBLISHABLE_KEY.",
            "Executar as migrações SQL no Supabase.",
            "Subir a base local pelo perfil Administrador.",
            "Fazer backup semanal em JSON.",
            "Manter o uso como piloto controlado até implementar Supabase Auth, RLS, auditoria e revisão LGPD.",
        ],
    )

    path = OUT / "Manual_Uso_SIDEP-CE_v0_7.docx"
    doc.save(path)
    return path


def build_tcc_update():
    doc = Document()
    style_doc(doc, "TCC SIDEP-CE - Atualização Metodológica v0.7")
    add_cover(
        doc,
        "SIDEP-CE: Atualização Metodológica v0.7",
        "Qualidade do banco de itens, exportação curricular e preparação pré-TRI/TRI",
        "0.7",
    )

    doc.add_heading("1. Síntese da Atualização", level=1)
    doc.add_paragraph(
        "A versão v0.7 consolida o SIDEP-CE como produto técnico-tecnológico aplicável em piloto controlado. "
        "A atualização fortalece a governança do banco de itens, a rastreabilidade das avaliações e a documentação operacional do sistema."
    )

    doc.add_heading("2. Relevância para o Mestrado", level=1)
    doc.add_paragraph(
        "As melhorias são relevantes para a pesquisa porque aproximam a proposta de uma aplicação real em contexto escolar. "
        "O sistema deixa de ser apenas protótipo demonstrativo e passa a possuir regras de negócio, documentação de uso, exportação da matriz avaliativa e preparação para coleta de dados empíricos."
    )

    doc.add_heading("3. Governança do Banco de Itens", level=1)
    doc.add_paragraph(
        "O banco de itens passou a ter controle contra duplicidade no cadastro de questões e contra repetição dentro da mesma prova. "
        "Essa regra preserva a diversidade de evidências, reduz risco de cola e melhora a validade pedagógica da avaliação por descritores."
    )

    doc.add_heading("4. Exportação Curricular", level=1)
    doc.add_paragraph(
        "A exportação em Markdown e PDF organiza curso, componente curricular, competência e descritor. "
        "Essa funcionalidade torna a matriz avaliativa auditável e facilita a validação por professores especialistas, coordenadores e banca acadêmica."
    )

    doc.add_heading("5. Relação com a Fase Pré-TRI", level=1)
    doc.add_paragraph(
        "O SIDEP-CE permanece em fase pré-TRI. A versão atual calcula indicadores diagnósticos por estudante, turma, descritor, componente e competência. "
        "A TRI plena exigirá volume maior de respostas, itens âncora, análise de dificuldade, discriminação e calibração psicométrica."
    )

    doc.add_heading("6. Produto Técnico-Tecnológico", level=1)
    add_bullets(
        doc,
        [
            "Plataforma online em React/Vite, Vercel e Supabase/PostgreSQL.",
            "Banco de itens estruturado por competências e descritores.",
            "Criação de avaliações com código único e imutável.",
            "Aplicação responsiva para computador, tablet e celular.",
            "Relatórios pedagógicos por aluno, turma, descritor, componente e competência.",
            "Manual operacional para uso por professores e gestores.",
        ],
    )

    doc.add_heading("7. Próxima Etapa Científica e Técnica", level=1)
    doc.add_paragraph(
        "A próxima etapa deve priorizar autenticação real, RLS por perfil, auditoria, LGPD e validação em campo. "
        "Do ponto de vista acadêmico, o próximo ciclo deve coletar evidências de uso em turmas reais, avaliar a qualidade dos itens e identificar descritores críticos recorrentes."
    )

    path = OUT / "TCC_SIDEP-CE_mestrado_metodologia_v0_7.docx"
    doc.save(path)
    return path


def main():
    manual = build_manual()
    tcc = build_tcc_update()
    print(manual)
    print(tcc)


if __name__ == "__main__":
    main()
