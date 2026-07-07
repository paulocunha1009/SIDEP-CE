from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\yolep\OneDrive\Documentos\Método TRI e IA")
OUT = ROOT / "outputs" / "TCC_SIDEP-CE_mestrado_metodologia_v0_6.docx"


TITLE = "SIDEP-CE: Sistema de Diagnóstico da Educação Profissional do Ceará baseado em Competências, TRI e Inteligência Artificial Generativa"
SUBTITLE = "Produto técnico-tecnológico para avaliação diagnóstica, formativa e final nos cursos técnicos da rede estadual"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in [
        ("Heading 1", 16, "0B3B75", 16, 8),
        ("Heading 2", 13, "0B3B75", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SIDEP-CE")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string("0B3B75")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TITLE)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(SUBTITLE)
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.italic = True

    doc.add_paragraph()
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    rows = [
        ("Autor", "Paulo Henrique Gomes Cordeiro da Cunha"),
        ("Área", "Educação Profissional e Tecnológica"),
        ("Curso piloto", "Técnico em Informática"),
        ("Versão", "v0.6 - MVP online em piloto controlado"),
        ("Data", "07 de julho de 2026"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        set_cell_text(row.cells[0], label, True)
        set_cell_text(row.cells[1], value)
        set_cell_shading(row.cells[0], "EAF2F8")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Ceará")
    run.font.size = Pt(11)

    doc.add_page_break()


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, True)
        set_cell_shading(cell, "EAF2F8")
    for values in rows:
        row = table.add_row()
        for cell, value in zip(row.cells, values):
            set_cell_text(cell, value)
    doc.add_paragraph()


def build():
    doc = Document()
    style_document(doc)
    add_title_page(doc)

    add_heading(doc, "Resumo", 1)
    add_para(
        doc,
        "Este trabalho apresenta o SIDEP-CE, Sistema de Diagnóstico da Educação Profissional do Ceará, como produto técnico-tecnológico voltado à avaliação diagnóstica, formativa e final dos cursos técnicos da rede estadual. A proposta articula matriz curricular, competências, descritores avaliáveis, banco de itens, análise pré-TRI, relatórios pedagógicos e uso supervisionado de Inteligência Artificial Generativa. O projeto parte da necessidade de transformar dados de desempenho técnico em intervenção pedagógica, apoiando professores, gestão escolar, CREDE/SEFOR e SEDUC na identificação de lacunas de aprendizagem e no planejamento de recomposição. Em sua versão v0.6, o SIDEP-CE evoluiu de protótipo local para MVP online em piloto controlado, com React/Vite, Vercel, Supabase/PostgreSQL, migração da base local, banco de itens de Informática e relatórios por aluno, turma, descritor, componente e competência.",
    )
    add_para(
        doc,
        "Palavras-chave: Educação Profissional; Avaliação Diagnóstica; Competências; Descritores; Teoria de Resposta ao Item; Learning Analytics; Inteligência Artificial Generativa.",
    )

    add_heading(doc, "1. Introdução", 1)
    add_para(
        doc,
        "A Educação Profissional e Tecnológica ocupa lugar estratégico no desenvolvimento social, produtivo e territorial do Ceará. Ao integrar formação geral, formação técnica, prática profissional, projetos integradores, laboratórios e preparação para o mundo do trabalho, os cursos técnicos exigem instrumentos avaliativos capazes de acompanhar aprendizagens específicas e orientar decisões pedagógicas em tempo oportuno.",
    )
    add_para(
        doc,
        "Nesse contexto, avaliações restritas à nota final ou a provas isoladas tendem a oferecer uma leitura limitada da aprendizagem técnica. O desafio central não é apenas verificar se o estudante acertou ou errou, mas compreender quais competências profissionais estão em desenvolvimento, quais descritores apresentam fragilidade, quais componentes curriculares exigem recomposição e como a escola pode agir para melhorar a qualidade da formação.",
    )
    add_para(
        doc,
        "O SIDEP-CE nasce como resposta a essa lacuna. A proposta estrutura um método e uma plataforma para diagnosticar, acompanhar e melhorar a aprendizagem técnica antes, durante e depois do curso, com foco inicial no curso Técnico em Informática e possibilidade de expansão para outros cursos e eixos tecnológicos.",
    )

    add_heading(doc, "2. Problema de Pesquisa", 1)
    add_para(
        doc,
        "Apesar da existência de matrizes curriculares e orientações normativas para a Educação Profissional, ainda é necessário operacionalizar um método contínuo de avaliação técnica que permita transformar a matriz em evidências observáveis de aprendizagem. O problema investigado pode ser sintetizado na seguinte questão: como estruturar um sistema de avaliação diagnóstica, formativa e final para cursos técnicos da rede estadual do Ceará, baseado em competências, descritores, banco de itens, análise pré-TRI e relatórios pedagógicos acionáveis?",
    )
    add_para(
        doc,
        "A ausência de dados sistemáticos por descritor dificulta a identificação de lacunas específicas. Sem esse nível de leitura, professores e gestores tendem a trabalhar com médias gerais, impressões isoladas ou resultados tardios. O SIDEP-CE propõe deslocar o foco da simples classificação para a tomada de decisão pedagógica.",
    )

    add_heading(doc, "3. Justificativa", 1)
    add_para(
        doc,
        "A relevância do projeto está em aproximar avaliação, currículo, gestão pedagógica e melhoria da aprendizagem técnica. Na Educação Profissional, o estudante precisa mobilizar conhecimentos, habilidades, atitudes e procedimentos em situações reais ou simuladas de trabalho. Portanto, a avaliação deve ser capaz de dialogar com componentes técnicos, laboratórios, projetos, estágio, tempo comunidade e práticas profissionais.",
    )
    add_para(
        doc,
        "Para o Ceará, o SIDEP-CE pode contribuir com uma base comum de diagnóstico da aprendizagem técnica, permitindo que escolas, regionais e gestão estadual acompanhem fragilidades por curso, componente, competência e descritor. Essa leitura apoia formação docente, produção de itens, planejamento de intervenções, revisão curricular e melhoria dos cursos ofertados.",
    )
    add_para(
        doc,
        "O projeto também possui valor acadêmico por tratar a TRI de forma metodologicamente responsável. O sistema não afirma realizar TRI plena na fase inicial. Ele estrutura uma base pré-TRI, com dificuldade inicial, histórico de respostas, curadoria de itens e relatórios diagnósticos, preparando o caminho para futura calibração psicométrica quando houver volume suficiente de dados.",
    )

    add_heading(doc, "4. Objetivos", 1)
    add_heading(doc, "4.1 Objetivo Geral", 2)
    add_para(
        doc,
        "Desenvolver e validar uma proposta metodológica e tecnológica para avaliação diagnóstica, formativa e final da Educação Profissional do Ceará, baseada em competências, descritores, banco de itens, análise pré-TRI, relatórios pedagógicos e uso supervisionado de IA Generativa.",
    )
    add_heading(doc, "4.2 Objetivos Específicos", 2)
    add_bullets(
        doc,
        [
            "Organizar uma matriz inicial de competências e descritores para o curso Técnico em Informática.",
            "Construir banco de itens vinculado a curso, componente curricular, competência e descritor.",
            "Criar fluxo de curadoria docente com status de rascunho, em revisão e validada.",
            "Implementar criador de avaliações com quantidade controlada de questões e uso exclusivo de itens validados.",
            "Permitir aplicação online por código único, com embaralhamento de questões e bloqueio de segunda tentativa.",
            "Gerar relatórios por aluno, turma, descritor, componente e competência.",
            "Estruturar base de dados para futura análise clássica de itens e calibração TRI.",
            "Documentar regras de negócio, governança, segurança e próximos passos para uso institucional.",
        ],
    )

    add_heading(doc, "5. Referencial Metodológico", 1)
    add_para(
        doc,
        "O SIDEP-CE se apoia em três eixos metodológicos: avaliação por competências, estruturação de descritores avaliáveis e análise progressiva dos dados de resposta. Competência é entendida como capacidade ampla de mobilizar conhecimentos, habilidades, atitudes e procedimentos em situações técnicas. Descritor é a evidência observável dessa competência, escrita de modo suficientemente específico para orientar itens, rubricas e relatórios.",
    )
    add_para(
        doc,
        "Na arquitetura avaliativa do sistema, a questão mede o descritor; o conjunto de descritores interpreta a competência; e os relatórios orientam a intervenção pedagógica. Essa separação é essencial porque a recomposição da aprendizagem precisa ocorrer no nível em que o professor consegue agir: descritor, componente curricular e prática formativa.",
    )
    add_para(
        doc,
        "A TRI aparece como horizonte psicométrico, não como promessa imediata. Na fase inicial, o sistema opera com indicadores diagnósticos, percentual de acerto, dificuldade inicial, histórico de respostas e análise por descritor. A calibração TRI requer volume de respostas, itens-âncora, estabilidade dos itens, análise de discriminação e validação por especialistas.",
    )

    add_heading(doc, "6. Produto Técnico-Tecnológico", 1)
    add_para(
        doc,
        "O produto desenvolvido é composto por dois elementos integrados: um caderno metodológico e uma plataforma digital. O caderno organiza fundamentos, matriz, regras de avaliação, governança e protocolo de evolução. A plataforma materializa esses princípios em um sistema web com áreas de aluno, professor, gestão escolar, regional, SEDUC e Administrador.",
    )
    add_table(
        doc,
        ["Camada", "Função no SIDEP-CE"],
        [
            ("Aluno", "Acessa avaliação por código e nome completo; não visualiza diagnóstico, peso de item ou relatório técnico."),
            ("Professor", "Cria banco de itens, valida questões, abre avaliações e consulta relatórios pedagógicos."),
            ("Gestão escolar", "Acompanha professores, avaliações e relatórios da escola."),
            ("CREDE/SEFOR", "Acompanha escolas da regional, com foco em indicadores agregados e apoio pedagógico."),
            ("SEDUC", "Acompanha visão de rede, cursos, eixos e políticas públicas."),
            ("Administrador", "Gerencia cadastros, permissões, senhas, banco de itens e parâmetros gerais."),
        ],
    )

    add_heading(doc, "7. Estado Atual do MVP v0.6", 1)
    add_para(
        doc,
        "Em 07 de julho de 2026, o SIDEP-CE encontra-se em fase de MVP online em piloto controlado. A aplicação foi implementada com React/Vite, versionada no GitHub, publicada via Vercel e conectada ao Supabase/PostgreSQL. O modo local foi preservado como fallback, permitindo continuidade do desenvolvimento e contingência.",
    )
    add_bullets(
        doc,
        [
            "Banco piloto do curso Técnico em Informática com 10 competências, 40 descritores e 441 questões.",
            "Tabelas online para competências, descritores, questões, avaliações, códigos bloqueados e respostas consolidadas.",
            "Rotina administrativa para migrar a base local do navegador para o Supabase.",
            "Código de avaliação randomizado, único e bloqueado após abertura ou exclusão.",
            "Aplicação do estudante com embaralhamento da ordem das questões.",
            "Respostas salvas em JSON consolidado por aluno/prova.",
            "Relatórios calculados sob demanda por aluno, turma, descritor, componente e competência.",
            "Backup semanal JSON disponível na área de relatórios.",
        ],
    )

    add_heading(doc, "8. Regras de Negócio Consolidadas", 1)
    add_bullets(
        doc,
        [
            "Toda questão deve estar vinculada a um descritor.",
            "Todo descritor deve estar vinculado a uma competência.",
            "Toda avaliação deve usar apenas questões validadas.",
            "A avaliação deve conter de 20 a 80 questões.",
            "O estudante só acessa avaliações com status aberta.",
            "O estudante não visualiza gabarito, peso, parâmetro pré-TRI, diagnóstico ou intervenção.",
            "Cada estudante pode responder apenas uma vez por código de avaliação e nome normalizado.",
            "A avaliação com respostas não pode ser excluída.",
            "O código de avaliação excluída permanece bloqueado e não pode ser reaproveitado.",
            "O uso online atual é piloto controlado até implantação de autenticação real, RLS e políticas de acesso.",
        ],
    )

    add_heading(doc, "9. Arquitetura Técnica", 1)
    add_para(
        doc,
        "A arquitetura atual foi escolhida para permitir avanço rápido, custo inicial zero e possibilidade de validação em campo. O React/Vite garante interface responsiva; a Vercel permite publicação web; o Supabase/PostgreSQL fornece banco online; e o GitHub mantém rastreabilidade do desenvolvimento.",
    )
    add_table(
        doc,
        ["Componente", "Tecnologia", "Papel"],
        [
            ("Interface", "React + Vite", "Aplicação web responsiva para aluno, professor e gestão."),
            ("Deploy", "Vercel", "Publicação online gratuita para piloto controlado."),
            ("Banco", "Supabase/PostgreSQL", "Persistência de cadastros, itens, avaliações e respostas."),
            ("Fallback", "localStorage", "Continuidade local e migração da base criada no protótipo."),
            ("Versionamento", "GitHub", "Histórico de código, documentação e deploy conectado."),
        ],
    )

    add_heading(doc, "10. Resultados Esperados", 1)
    add_para(
        doc,
        "Espera-se que o SIDEP-CE permita identificar, com maior precisão, os descritores técnicos em que os estudantes apresentam maior dificuldade. Essa informação deve orientar intervenções pedagógicas, formação docente, produção de novos itens e planejamento institucional.",
    )
    add_bullets(
        doc,
        [
            "Diagnóstico inicial do nível técnico dos estudantes.",
            "Acompanhamento formativo durante o curso.",
            "Avaliação final com leitura por competência e descritor.",
            "Identificação de descritores críticos por turma e escola.",
            "Apoio ao planejamento de recomposição das aprendizagens.",
            "Construção de base histórica para análise estatística e futura TRI.",
            "Fortalecimento da gestão pedagógica da Educação Profissional.",
        ],
    )

    add_heading(doc, "11. Discussão", 1)
    add_para(
        doc,
        "A principal contribuição do SIDEP-CE é deslocar a avaliação técnica de uma lógica pontual para uma lógica sistêmica, contínua e pedagógica. Ao organizar competências e descritores, o sistema permite que o professor interprete o erro não como falha genérica, mas como evidência de uma habilidade específica que pode ser retomada.",
    )
    add_para(
        doc,
        "A implantação online amplia o potencial do projeto porque permite coleta de dados entre turmas e escolas, criando uma base comum para análise. Entretanto, essa ampliação exige responsabilidade: autenticação institucional, RLS, auditoria, LGPD e política clara de uso dos dados são condições para expansão segura.",
    )
    add_para(
        doc,
        "O uso de IA Generativa deve permanecer supervisionado. A IA pode apoiar a criação de trilhas, revisão de itens, síntese de relatórios e sugestão de atividades, mas não deve substituir o professor, decidir nota, classificar estudante ou acessar dados pessoais sem governança.",
    )

    add_heading(doc, "12. Plano de Evolução", 1)
    add_table(
        doc,
        ["Fase", "Entregas Prioritárias"],
        [
            ("Piloto controlado", "Aplicar com turmas pequenas, validar fluxo, conferir registros no Supabase e ajustar usabilidade."),
            ("Segurança", "Implementar Supabase Auth, RLS, políticas por perfil, logs e recuperação segura de senha."),
            ("Gestão escolar", "Ampliar painéis por escola, turma, curso, descritor e intervenção pedagógica."),
            ("Análise estatística", "Calcular dificuldade empírica, discriminação, distratores e itens candidatos a âncora."),
            ("TRI", "Calibrar itens após volume suficiente de respostas e validação psicométrica."),
            ("IA supervisionada", "Gerar trilhas e feedbacks sob revisão humana, com anonimização e registro de uso."),
        ],
    )

    add_heading(doc, "13. Considerações Finais", 1)
    add_para(
        doc,
        "O SIDEP-CE apresenta uma proposta metodológica e tecnológica viável para fortalecer a avaliação da Educação Profissional do Ceará. A passagem do protótipo local para MVP online representa um avanço importante, pois permite testar a metodologia em ambiente real, coletar respostas, gerar relatórios e amadurecer a base de dados necessária à futura calibração TRI.",
    )
    add_para(
        doc,
        "Sua relevância não está apenas na tecnologia, mas na capacidade de transformar dados em intervenção pedagógica. Ao vincular avaliação a competências, descritores, banco de itens e relatórios acionáveis, o sistema cria condições para que professores e gestores compreendam melhor a aprendizagem técnica e atuem de forma mais precisa.",
    )
    add_para(
        doc,
        "Como produto de mestrado, o SIDEP-CE articula pesquisa aplicada, inovação pedagógica, desenvolvimento tecnológico e compromisso com a qualidade da Educação Profissional. A continuidade do projeto deve priorizar segurança, validação em campo, análise estatística dos itens e expansão gradual para outros cursos técnicos da rede estadual.",
    )

    add_heading(doc, "Referências e Norteadores", 1)
    add_bullets(
        doc,
        [
            "Matrizes curriculares do curso Técnico em Informática utilizadas no projeto SIDEP-CE.",
            "Catálogo Nacional de Cursos Técnicos como referência para eixo tecnológico e perfil profissional.",
            "Documentos normativos e curriculares informados pelo autor do projeto para alinhamento à Educação Profissional do Ceará.",
            "Artigo, resumo, banner e TCC SIDEP-CE produzidos no ciclo do Seminário DoCEntes 2026.",
            "Documentação técnica do MVP SIDEP-CE, incluindo schema PostgreSQL, backlog Scrum/MCP e documentação de implantação online gratuita.",
        ],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT)
