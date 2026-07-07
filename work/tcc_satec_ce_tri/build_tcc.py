from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

OUT = Path(r"C:\Users\yolep\OneDrive\Documentos\Método TRI e IA\outputs\TCC_Mestrado_SATEC_CE_TRI.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(9.5)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.05
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, True)
        set_cell_shading(hdr[i], "E8EEF5")
        if widths:
            hdr[i].width = widths[i]
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
            if widths:
                cells[i].width = widths[i]
    doc.add_paragraph()
    return table


def add_para(doc, text, style=None, align=None):
    p = doc.add_paragraph(style=style)
    p.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.35) if style is None else None
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        p.add_run(item)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


doc = Document()

section = doc.sections[0]
section.page_height = Inches(11)
section.page_width = Inches(8.5)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.49)
section.footer_distance = Inches(0.49)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)
styles["Normal"].font.color.rgb = RGBColor(31, 41, 55)
styles["Normal"].paragraph_format.space_after = Pt(6)
styles["Normal"].paragraph_format.line_spacing = 1.15
for name, size, color in [
    ("Heading 1", 16, RGBColor(46, 116, 181)),
    ("Heading 2", 13, RGBColor(46, 116, 181)),
    ("Heading 3", 12, RGBColor(31, 77, 120)),
]:
    styles[name].font.name = "Calibri"
    styles[name].font.size = Pt(size)
    styles[name].font.color.rgb = color
    styles[name].font.bold = True

footer = section.footer.paragraphs[0]
footer.text = "SATEC-CE/TRI - Projeto de pesquisa em nível de mestrado profissional"
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.runs[0].font.size = Pt(8)
footer.runs[0].font.color.rgb = RGBColor(107, 114, 128)

# Capa
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("INSTITUIÇÃO DE ENSINO SUPERIOR\nPROGRAMA DE PÓS-GRADUAÇÃO EM EDUCAÇÃO PROFISSIONAL E TECNOLÓGICA")
r.bold = True
r.font.size = Pt(12)
doc.add_paragraph("\n\n")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("SATEC-CE/TRI: SISTEMA DE AVALIAÇÃO TÉCNICA DA EDUCAÇÃO PROFISSIONAL DO CEARÁ BASEADO NA TEORIA DE RESPOSTA AO ITEM")
r.bold = True
r.font.size = Pt(15)
doc.add_paragraph("\n")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Projeto de pesquisa e produto técnico-tecnológico em nível de mestrado profissional")
r.italic = True
r.font.size = Pt(12)
doc.add_paragraph("\n\n\n")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Autor: [Nome do pesquisador]\nOrientador(a): [Nome do(a) orientador(a)]")
doc.add_paragraph("\n\n\n")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Fortaleza - CE\n2026")
doc.add_page_break()

# Resumo
heading(doc, "RESUMO", 1)
add_para(
    doc,
    "Este trabalho apresenta uma proposta de desenvolvimento e validação do SATEC-CE/TRI, um método de avaliação diagnóstica, formativa e final para cursos técnicos das escolas profissionalizantes da rede estadual do Ceará. A proposta articula Educação Profissional e Tecnológica, Documento Curricular Referencial do Ceará, matrizes curriculares das Escolas Estaduais de Educação Profissional, Catálogo Nacional de Cursos Técnicos, avaliação por competências, Learning Analytics, Inteligência Artificial Generativa e Teoria de Resposta ao Item. Parte-se do problema de que os cursos técnicos possuem matrizes curriculares extensas e bem estruturadas, mas ainda carecem de um sistema contínuo de avaliação capaz de produzir evidências comparáveis sobre o desenvolvimento das competências profissionais dos estudantes antes, durante e ao final do percurso formativo. O objetivo geral é desenvolver e validar um método de avaliação técnica que transforme componentes curriculares e perfis profissionais de conclusão em descritores avaliáveis, banco de itens, rubricas práticas, escala de proficiência e relatórios pedagógicos para tomada de decisão. A metodologia proposta é aplicada, de abordagem mista e inspiração em pesquisa-intervenção e design-based research, com estudo de caso em cursos-piloto da rede estadual. O produto educacional previsto compreende um caderno metodológico, matrizes avaliativas por curso, banco inicial de itens, rubricas de desempenho, protocolo de aplicação e modelo de dashboard diagnóstico. Espera-se que a proposta contribua para a recomposição das aprendizagens, o acompanhamento longitudinal dos estudantes, a melhoria dos cursos técnicos e a formulação de estratégias de gestão pedagógica baseadas em evidências.",
)
add_para(doc, "Palavras-chave: Educação Profissional e Tecnológica. Teoria de Resposta ao Item. Avaliação por competências. Ceará. Learning Analytics. Inteligência Artificial Generativa.", style=None)

heading(doc, "ABSTRACT", 1)
add_para(
    doc,
    "This study presents a proposal for the development and validation of SATEC-CE/TRI, a diagnostic, formative and final assessment method for technical courses offered by the public professional education network of Ceará, Brazil. The proposal connects Professional and Technological Education, state curriculum guidelines, technical course matrices, competency-based assessment, Learning Analytics, Generative Artificial Intelligence and Item Response Theory. The main problem is that technical courses have robust curricular structures but still lack a continuous assessment system capable of producing comparable evidence about students' professional competence development before, during and after the course. The general objective is to develop and validate a technical assessment method that converts curricular components and professional profiles into measurable descriptors, item banks, practical rubrics, proficiency scales and pedagogical reports. The methodology is applied, mixed-methods and inspired by intervention research and design-based research, with case studies in pilot technical courses. The expected educational product includes a methodological guide, course assessment matrices, an initial item bank, performance rubrics, an application protocol and a diagnostic dashboard model.",
)
add_para(doc, "Keywords: Professional and Technological Education. Item Response Theory. Competency-based assessment. Ceará. Learning Analytics. Generative Artificial Intelligence.")
doc.add_page_break()

# Sumário manual
heading(doc, "SUMÁRIO", 1)
toc = [
    "1 INTRODUÇÃO",
    "2 PROBLEMA DE PESQUISA",
    "3 JUSTIFICATIVA",
    "4 OBJETIVOS",
    "5 REFERENCIAL TEÓRICO",
    "6 BASE CURRICULAR E NORMATIVA",
    "7 PROPOSTA DO MÉTODO SATEC-CE/TRI",
    "8 METODOLOGIA DA PESQUISA",
    "9 PRODUTO TÉCNICO-TECNOLÓGICO",
    "10 CRONOGRAMA",
    "11 RESULTADOS ESPERADOS",
    "12 CONSIDERAÇÕES FINAIS",
    "REFERÊNCIAS",
    "APÊNDICES",
]
for item in toc:
    add_para(doc, item, align=WD_ALIGN_PARAGRAPH.LEFT)
doc.add_page_break()

heading(doc, "1 INTRODUÇÃO", 1)
add_para(
    doc,
    "A Educação Profissional e Tecnológica ocupa posição estratégica no desenvolvimento educacional, econômico e social do Ceará. As Escolas Estaduais de Educação Profissional consolidaram, ao longo dos últimos anos, uma política pública voltada à integração entre Ensino Médio e formação técnica, articulando formação geral, formação profissional, parte diversificada, estágio curricular e aproximação com o mundo do trabalho. O Caderno de Matrizes 2026 da Secretaria da Educação do Ceará registra a oferta de cinquenta e dois cursos técnicos nas EEEPs, distribuídos em eixos tecnológicos como Ambiente e Saúde, Controle e Processos Industriais, Gestão e Negócios, Informação e Comunicação, Infraestrutura, Produção Alimentícia, Produção Cultural e Design, Produção Industrial, Recursos Naturais, Segurança e Turismo, Hospitalidade e Lazer.",
)
add_para(
    doc,
    "A existência de matrizes curriculares extensas e organizadas, entretanto, não garante por si só a produção de evidências contínuas sobre o desenvolvimento das competências dos estudantes. Em muitas experiências escolares, a avaliação permanece restrita à nota disciplinar, ao percentual de acertos ou ao cumprimento formal da carga horária. Essa forma de acompanhamento é importante, mas insuficiente para responder a perguntas decisivas para a gestão pedagógica da EPT: quais competências técnicas os estudantes dominam ao ingressar no curso? Em quais descritores há maior fragilidade ao longo do percurso? Que habilidades devem ser priorizadas antes do estágio? Como comparar a evolução de turmas, escolas, cursos e eixos tecnológicos sem reduzir a análise a médias brutas?",
)
add_para(
    doc,
    "A Teoria de Resposta ao Item oferece uma alternativa metodológica para enfrentar esse problema. Diferentemente da simples contagem de acertos, a TRI estima proficiência considerando as características dos itens e o padrão de respostas dos estudantes. Essa lógica é amplamente utilizada em avaliações educacionais de larga escala, como Saeb, Enem e SPAECE, permitindo construir escalas comparáveis, acompanhar evolução e interpretar resultados por níveis de desempenho. O desafio desta pesquisa é adaptar esse racional psicométrico ao contexto dos cursos técnicos, em que a aprendizagem envolve não apenas conhecimentos conceituais, mas também competências práticas, tomada de decisão, resolução de problemas, postura profissional e domínio de procedimentos.",
)
add_para(
    doc,
    "Nesse contexto, propõe-se o SATEC-CE/TRI, Sistema de Avaliação Técnica da Educação Profissional do Ceará baseado na Teoria de Resposta ao Item. O método pretende transformar matrizes curriculares e perfis profissionais de conclusão em descritores avaliáveis, banco de itens, rubricas práticas e relatórios de proficiência. A proposta dialoga com a Avaliação da Educação Profissional e Tecnológica desenvolvida pelo Inep, especialmente com a noção de Sistema Nacional de Avaliação da EPT, que considera dimensões como articulação com o mundo do trabalho, condições de oferta, permanência e conclusão, desenvolvimento dos estudantes e impacto social da formação profissional.",
)

heading(doc, "2 PROBLEMA DE PESQUISA", 1)
add_para(
    doc,
    "O problema central desta pesquisa pode ser formulado nos seguintes termos: como desenvolver e validar um método de avaliação diagnóstica, formativa e final, baseado em competências e inspirado nos parâmetros da Teoria de Resposta ao Item, capaz de acompanhar a proficiência técnica dos estudantes dos cursos profissionalizantes da rede estadual do Ceará antes, durante e ao final do percurso formativo?",
)
add_para(
    doc,
    "O problema deriva de uma tensão entre a complexidade curricular da EPT e a limitação dos instrumentos avaliativos usuais. Os cursos técnicos integrados ao Ensino Médio são compostos por múltiplas dimensões: formação geral básica, formação técnica, práticas laboratoriais, estágio, projetos integradores e experiências ligadas ao território. Todavia, a avaliação escolar nem sempre consegue traduzir esse conjunto em indicadores claros de progressão de competências. Em consequência, professores e gestores podem até identificar que uma turma apresenta baixo rendimento, mas têm mais dificuldade de localizar quais habilidades específicas exigem recomposição, aprofundamento ou reorganização didática.",
)
add_para(
    doc,
    "A questão norteadora da pesquisa é: de que modo um sistema de avaliação técnica baseado em matriz de competências, descritores, itens calibráveis e rubricas práticas pode apoiar a melhoria dos cursos técnicos das escolas profissionalizantes do Ceará, gerando dados úteis para estudantes, professores, coordenações escolares e gestão estadual?",
)

heading(doc, "3 JUSTIFICATIVA", 1)
add_para(
    doc,
    "A justificativa acadêmica da pesquisa está no encontro entre três campos que frequentemente são estudados separadamente: avaliação educacional, Educação Profissional e Tecnológica e tecnologias de análise de dados educacionais. Embora a TRI seja consolidada nas avaliações de larga escala da educação básica, sua aplicação sistemática à avaliação de competências técnicas em redes estaduais ainda demanda investigação, adaptação metodológica e validação empírica. A proposta busca contribuir para esse campo ao propor um modelo que respeita a especificidade da EPT, sem importar de forma mecânica a lógica de avaliações exclusivamente cognitivas.",
)
add_para(
    doc,
    "A justificativa pedagógica está na necessidade de transformar a avaliação em instrumento de aprendizagem. Um sistema de avaliação técnica deve permitir que o estudante compreenda seu próprio percurso, identifique competências já consolidadas e reconheça pontos de melhoria. Para o professor, os resultados devem indicar descritores frágeis, grupos de recomposição e estratégias didáticas específicas. Para a escola, os relatórios devem favorecer planejamento por curso, por eixo tecnológico e por série. Para a rede, a avaliação deve gerar metas e estratégias de melhoria do ensino profissionalizante, com base em evidências comparáveis.",
)
add_para(
    doc,
    "A justificativa social e institucional decorre do papel das escolas profissionalizantes no projeto de desenvolvimento do Ceará. A qualidade da formação técnica impacta a inserção produtiva dos jovens, a continuidade dos estudos, a inovação local e a capacidade das escolas responderem às demandas dos territórios. Ao propor um método de avaliação antes, durante e depois do curso, esta pesquisa busca criar condições para que a rede não apenas certifique estudantes, mas acompanhe a construção efetiva do perfil profissional de conclusão.",
)
add_para(
    doc,
    "A proposta também se justifica pela existência de documentos e sistemas nacionais que podem ser apropriados e aprimorados na realidade cearense. O Inep já organiza a Avaliação da Educação Profissional e Tecnológica em dimensões articuladas, e as matrizes curriculares da SEDUC-CE já descrevem a estrutura dos cursos ofertados. O SATEC-CE/TRI pretende atuar como ponte entre esses referenciais e a prática cotidiana das escolas.",
)

heading(doc, "4 OBJETIVOS", 1)
heading(doc, "4.1 Objetivo geral", 2)
add_para(
    doc,
    "Desenvolver e validar uma proposta metodológica de avaliação diagnóstica, formativa e final para cursos técnicos da rede estadual do Ceará, baseada em matriz de competências, descritores avaliáveis, banco de itens, rubricas práticas e parâmetros inspirados na Teoria de Resposta ao Item.",
)
heading(doc, "4.2 Objetivos específicos", 2)
add_bullets(
    doc,
    [
        "Mapear os cursos técnicos ofertados pela rede estadual do Ceará e seus eixos tecnológicos, a partir das matrizes curriculares 2025/2026.",
        "Construir uma matriz avaliativa piloto que relacione componentes curriculares, competências, descritores e evidências de aprendizagem.",
        "Elaborar instrumentos de avaliação para três momentos: entrada, percurso e saída do curso técnico.",
        "Definir níveis de proficiência técnica que permitam interpretar o desenvolvimento dos estudantes em escala progressiva.",
        "Propor um banco inicial de itens e rubricas práticas passíveis de calibração psicométrica.",
        "Desenvolver um modelo de devolutiva pedagógica com relatórios para estudante, professor, coordenação e gestão.",
        "Investigar o uso de Inteligência Artificial Generativa como apoio à elaboração de trilhas de recomposição, atividades personalizadas e análise pedagógica dos resultados.",
    ],
)

heading(doc, "5 REFERENCIAL TEÓRICO", 1)
heading(doc, "5.1 Avaliação educacional e Teoria de Resposta ao Item", 2)
add_para(
    doc,
    "A avaliação educacional pode ser compreendida como processo de produção de evidências para orientar decisões pedagógicas, curriculares e institucionais. Quando reduzida à atribuição de notas, perde parte de sua potência formativa. Quando organizada em torno de competências, descritores e devolutivas, torna-se instrumento de diagnóstico e intervenção. A Teoria de Resposta ao Item, nesse cenário, oferece contribuição relevante por permitir estimar proficiência com base no padrão de respostas e nas propriedades dos itens.",
)
add_para(
    doc,
    "Nos modelos clássicos da TRI, cada item pode ser descrito por parâmetros como dificuldade, discriminação e, em modelos de três parâmetros, probabilidade de acerto casual. Essa abordagem permite construir escalas em que os estudantes são posicionados não apenas pelo número de acertos, mas pelo tipo de item que conseguiram resolver. Em avaliações como Saeb e Enem, esse princípio permite comparabilidade entre edições e interpretação pedagógica por níveis de desempenho.",
)
add_para(
    doc,
    "A adaptação da TRI à EPT exige cuidado metodológico. Cursos técnicos envolvem tarefas práticas, protocolos, atitudes profissionais e resolução de situações reais ou simuladas. Por isso, o SATEC-CE/TRI não propõe substituir toda avaliação por prova objetiva. A proposta é articular itens calibráveis, tarefas práticas e rubricas de desempenho, compondo um sistema híbrido de evidências. Em uma fase inicial, a pesquisa pode operar com estatística descritiva e análise clássica dos itens; à medida que o banco de respostas cresce, torna-se possível calibrar itens e construir escalas de proficiência técnica.",
)

heading(doc, "5.2 Avaliação por competências na Educação Profissional", 2)
add_para(
    doc,
    "A Educação Profissional e Tecnológica não se limita à transmissão de conteúdos. Seu compromisso é formar sujeitos capazes de mobilizar conhecimentos, habilidades, atitudes e valores em situações de trabalho e de vida social. A avaliação por competências, portanto, deve observar a capacidade do estudante de interpretar problemas, escolher procedimentos, executar tarefas, justificar decisões, comunicar resultados e atuar com responsabilidade técnica, ética e segurança.",
)
add_para(
    doc,
    "No contexto dos cursos técnicos integrados, a avaliação por competências precisa dialogar com a formação geral. A leitura técnica, a matemática aplicada, a interpretação de gráficos, a comunicação profissional e o raciocínio lógico são condições para o desempenho técnico. Assim, uma matriz avaliativa da EPT deve integrar competências gerais e competências profissionais, evitando a fragmentação entre disciplinas da base comum e componentes técnicos.",
)

heading(doc, "5.3 Learning Analytics e tomada de decisão pedagógica", 2)
add_para(
    doc,
    "Learning Analytics refere-se à mensuração, coleta, análise e apresentação de dados sobre estudantes e seus contextos, com a finalidade de compreender e otimizar a aprendizagem. Aplicado à EPT, esse campo permite transformar resultados avaliativos em informação acionável para professores e gestores. A proposta do SATEC-CE/TRI aproxima Learning Analytics e avaliação por competências ao defender que cada resultado deve gerar uma consequência pedagógica: recomposição, aprofundamento, reagrupamento, intervenção prática ou revisão curricular.",
)
add_para(
    doc,
    "Essa perspectiva é especialmente importante na avaliação durante o curso. Se a avaliação ocorre apenas ao final, sua capacidade de melhorar o percurso é limitada. Ao organizar diagnósticos de entrada, avaliações de percurso e avaliação de saída, o método permite identificar lacunas antes que se consolidem e acompanhar a progressão do estudante ao longo dos semestres.",
)

heading(doc, "5.4 Inteligência Artificial Generativa como apoio à intervenção", 2)
add_para(
    doc,
    "A Inteligência Artificial Generativa pode apoiar a personalização da aprendizagem, desde que inserida em um desenho pedagógico responsável. No SATEC-CE/TRI, a IA não ocupa o lugar do professor nem define sozinha o julgamento avaliativo. Sua função é apoiar a produção de atividades por descritor, trilhas de recomposição, simulados equivalentes, feedbacks formativos e relatórios preliminares. A decisão pedagógica permanece com docentes e equipes escolares.",
)
add_para(
    doc,
    "O uso de IA exige governança ética dos dados. A pesquisa deve observar proteção de dados pessoais, minimização de informações sensíveis, transparência no uso dos resultados e revisão humana das recomendações geradas. Em contexto escolar, a IA deve ser instrumento de apoio à equidade e não mecanismo de rotulação permanente dos estudantes.",
)

heading(doc, "6 BASE CURRICULAR E NORMATIVA", 1)
add_para(
    doc,
    "A proposta se apoia em documentos curriculares, legais e institucionais que organizam a Educação Profissional e o Ensino Médio no Brasil e no Ceará. Entre eles, destacam-se a Lei de Diretrizes e Bases da Educação Nacional, a Base Nacional Comum Curricular, o Documento Curricular Referencial do Ceará, o Catálogo Nacional de Cursos Técnicos, as Diretrizes Curriculares Nacionais Gerais para a Educação Profissional e Tecnológica, a Lei nº 14.945/2024 e as matrizes curriculares da SEDUC-CE para 2025/2026.",
)
add_para(
    doc,
    "O Caderno de Matrizes 2026 informa que a formação geral básica nas turmas ingressantes passou a atender à Lei nº 14.945/2024, com carga horária de 2.880 h/a, correspondente às 2.400 horas mínimas. Também informa que a composição das matrizes das EEEPs contempla formação geral, formação profissional, parte diversificada e estágio curricular, totalizando 5.400 h/a ao longo das três séries do Ensino Médio integrado à Educação Profissional.",
)
add_para(
    doc,
    "Os documentos das Escolas do Campo e da Escola Família Agrícola evidenciam outra dimensão importante: a necessidade de considerar pedagogia da alternância, tempo escola, tempo comunidade, práticas sociais comunitárias, organização do trabalho e técnicas produtivas. Portanto, o método proposto deve ser flexível o bastante para avaliar cursos técnicos urbanos, cursos do campo, ofertas integradas, subsequentes e experiências com forte territorialidade.",
)
add_para(
    doc,
    "O proponente do projeto indicou ainda a necessidade de alinhamento com PNE 15388/2026, LC 212/2025, Lei 14.945/2024, Resolução CNE/CEB 4/2025, Decreto 12.603/2025, Lei 12.603/2025, CNCT e educação híbrida. No texto final da pesquisa, recomenda-se que essas normas sejam conferidas em versão oficial e incorporadas em quadro normativo próprio, distinguindo legislação federal, estadual, atos do CNE, decretos regulamentadores e documentos curriculares da SEDUC-CE. Essa cautela é necessária para garantir precisão jurídica e validade institucional da proposta.",
)

add_table(
    doc,
    ["Referencial", "Contribuição para o SATEC-CE/TRI"],
    [
        ["DCRC", "Orienta a formação integral, o currículo do Ensino Médio no Ceará e a articulação entre competências, território e projeto de vida."],
        ["Matrizes 2025/2026 da SEDUC-CE", "Definem cursos ofertados, carga horária, componentes curriculares, séries e organização da formação profissional."],
        ["CNCT", "Define eixos tecnológicos, perfis profissionais de conclusão, carga horária mínima e identidade nacional dos cursos técnicos."],
        ["Sinaept/Inep", "Oferece referência para dimensões de avaliação da EPT: mundo do trabalho, condições de oferta, permanência, desenvolvimento e impacto social."],
        ["TRI", "Permite construir escala de proficiência e interpretar resultados para além do percentual bruto de acertos."],
        ["Educação híbrida", "Exige diferenciar evidências presenciais, remotas, laboratoriais, simuladas, supervisionadas e territoriais."],
    ],
)

heading(doc, "7 PROPOSTA DO MÉTODO SATEC-CE/TRI", 1)
add_para(
    doc,
    "O SATEC-CE/TRI é concebido como um método de avaliação técnica organizado em ciclos. Cada ciclo parte da matriz curricular do curso, identifica competências, constrói descritores, elabora instrumentos, aplica avaliações, analisa resultados e gera intervenções pedagógicas. O método pretende ser escalável, mas deve iniciar por cursos-piloto para validação conceitual e psicométrica.",
)
heading(doc, "7.1 Dimensões avaliativas", 2)
add_table(
    doc,
    ["Dimensão", "O que avalia", "Exemplos de evidência"],
    [
        ["Formação geral aplicada", "Leitura técnica, matemática aplicada, interpretação de normas, gráficos e problemas.", "Itens objetivos, estudos de caso, problemas contextualizados."],
        ["Competências técnicas específicas", "Conhecimentos e habilidades próprias do curso e do eixo tecnológico.", "Itens técnicos, simulações, tarefas de laboratório."],
        ["Prática profissional", "Execução de procedimentos, uso de equipamentos, segurança e qualidade.", "Rubricas, checklists, protocolos práticos."],
        ["Mundo do trabalho", "Ética, comunicação, organização produtiva, empreendedorismo e trabalho em equipe.", "Projetos, relatórios, apresentação técnica."],
        ["Condições de oferta", "Infraestrutura, laboratório, conectividade, insumos e suporte pedagógico.", "Instrumento institucional, questionários, observação."],
        ["Permanência e progressão", "Evolução por semestre, frequência, risco de evasão e recuperação.", "Dashboard, histórico escolar, acompanhamento docente."],
    ],
)

heading(doc, "7.2 Momentos de aplicação", 2)
add_para(
    doc,
    "A avaliação de entrada deve ocorrer no início do curso ou da série, com foco em pré-requisitos e conhecimentos iniciais. Nos cursos de Informação e Comunicação, por exemplo, pode avaliar raciocínio lógico, leitura de problemas, matemática básica, cultura digital e noções de hardware ou software. Nos cursos de Recursos Naturais, pode avaliar leitura de território, fundamentos de biologia, matemática aplicada e conhecimentos sobre práticas produtivas.",
)
add_para(
    doc,
    "A avaliação de percurso deve ocorrer ao final de módulos, semestres ou unidades curriculares estratégicas. Seu objetivo não é classificar estudantes, mas orientar intervenção. Os resultados devem indicar descritores frágeis, grupos de recomposição, necessidades de laboratório, revisão de metodologias e reorganização das atividades práticas.",
)
add_para(
    doc,
    "A avaliação de saída deve ocorrer antes da conclusão do curso, estágio final ou apresentação de projeto integrador. Deve verificar se o estudante se aproxima do perfil profissional de conclusão previsto no CNCT e na matriz curricular. Nessa etapa, recomenda-se combinar prova técnica, situação-problema, rubrica prática e evidência de projeto ou estágio.",
)

heading(doc, "7.3 Níveis de proficiência técnica", 2)
add_table(
    doc,
    ["Nível", "Descrição sintética", "Interpretação pedagógica"],
    [
        ["Nível 1 - Inicial", "Reconhece conceitos básicos, mas depende de orientação constante.", "Requer recomposição de pré-requisitos e acompanhamento próximo."],
        ["Nível 2 - Básico operacional", "Executa procedimentos simples em situações conhecidas.", "Precisa ampliar repertório e praticar variações de tarefa."],
        ["Nível 3 - Intermediário", "Resolve problemas técnicos previsíveis e justifica procedimentos.", "Pode avançar para situações integradoras e projetos."],
        ["Nível 4 - Autônomo", "Seleciona estratégias, executa com segurança e adapta procedimentos.", "Está próximo do perfil profissional esperado."],
        ["Nível 5 - Integrador", "Integra conhecimentos, propõe soluções e comunica decisões técnicas.", "Demonstra prontidão para desafios complexos, estágio e inovação."],
    ],
)

heading(doc, "8 METODOLOGIA DA PESQUISA", 1)
add_para(
    doc,
    "A pesquisa caracteriza-se como aplicada, de abordagem mista, com natureza exploratória e interventiva. A escolha por uma abordagem mista justifica-se pela necessidade de combinar dados quantitativos, como desempenho em itens e evolução de proficiência, com dados qualitativos, como percepção docente, análise de tarefas práticas e adequação dos instrumentos ao contexto real dos cursos técnicos.",
)
add_para(
    doc,
    "O desenho metodológico aproxima-se da pesquisa-intervenção e da design-based research, pois busca desenvolver um produto educacional, aplicá-lo em contexto real, coletar evidências, revisá-lo e produzir conhecimento sobre sua implementação. O estudo de caso será realizado em cursos-piloto da rede estadual, preferencialmente contemplando eixos tecnológicos distintos para testar a flexibilidade do método.",
)
heading(doc, "8.1 Campo e participantes", 2)
add_para(
    doc,
    "O campo de pesquisa será constituído por escolas profissionalizantes da rede estadual do Ceará. Recomenda-se iniciar com cinco cursos-piloto: Técnico em Informática ou Desenvolvimento de Sistemas, Técnico em Administração, Técnico em Agroecologia ou Agropecuária, Técnico em Enfermagem e Técnico em Eletrotécnica ou Automação Industrial. Essa seleção contempla tecnologia, gestão, campo, saúde e indústria, permitindo observar diferentes tipos de competência e evidência prática.",
)
add_para(
    doc,
    "Os participantes poderão incluir estudantes, professores da formação técnica, professores da formação geral, coordenadores de curso, gestores escolares e técnicos da rede. A participação em coleta de dados deverá observar autorização institucional, consentimento e, quando aplicável, submissão a comitê de ética em pesquisa.",
)
heading(doc, "8.2 Etapas metodológicas", 2)
add_numbered(
    doc,
    [
        "Análise documental das matrizes curriculares, DCRC, CNCT e documentos de avaliação da EPT.",
        "Seleção dos cursos-piloto e definição das competências prioritárias.",
        "Construção da matriz avaliativa por curso, com descritores e evidências esperadas.",
        "Elaboração de banco inicial de itens, situações-problema e rubricas práticas.",
        "Validação de conteúdo por professores especialistas e coordenadores técnicos.",
        "Aplicação da avaliação diagnóstica de entrada.",
        "Análise dos resultados por descritor, estudante, turma e curso.",
        "Planejamento de intervenção pedagógica com apoio de IA Generativa supervisionada.",
        "Aplicação de avaliação de percurso e ajuste dos instrumentos.",
        "Aplicação de avaliação de saída e elaboração de relatório final.",
    ],
)
heading(doc, "8.3 Instrumentos de coleta", 2)
add_table(
    doc,
    ["Instrumento", "Finalidade"],
    [
        ["Matriz avaliativa", "Relacionar componentes curriculares, competências, descritores e tipos de evidência."],
        ["Teste diagnóstico", "Medir conhecimentos de entrada e pré-requisitos."],
        ["Teste de percurso", "Acompanhar desenvolvimento semestral de competências."],
        ["Situação-problema", "Avaliar aplicação técnica em contexto próximo do mundo do trabalho."],
        ["Rubrica prática", "Registrar desempenho em laboratório, campo, estágio ou projeto."],
        ["Questionário docente", "Coletar percepção sobre validade, clareza e utilidade dos resultados."],
        ["Dashboard diagnóstico", "Organizar devolutivas por estudante, turma, curso, escola e eixo."],
    ],
)
heading(doc, "8.4 Análise dos dados", 2)
add_para(
    doc,
    "Na fase inicial, os dados serão analisados por estatística descritiva, percentual de acerto por descritor, distribuição de desempenho, consistência dos itens e comparação entre avaliação de entrada e avaliação de percurso. Quando houver volume suficiente de respostas, será possível avançar para calibração de itens com modelos de TRI, inicialmente considerando itens dicotômicos e, posteriormente, modelos adequados a itens politômicos ou rubricas graduadas.",
)
add_para(
    doc,
    "Os dados qualitativos serão analisados por categorização temática, considerando percepções docentes sobre clareza dos descritores, pertinência dos itens, adequação das rubricas e utilidade pedagógica dos relatórios. A triangulação entre dados quantitativos e qualitativos permitirá revisar o método e aprimorar o produto educacional.",
)

heading(doc, "8.5 Estratégia de validação psicométrica", 2)
add_para(
    doc,
    "A validação psicométrica do SATEC-CE/TRI será organizada em etapas progressivas, respeitando o fato de que uma escala baseada em TRI não nasce pronta. Na primeira aplicação, a análise deverá observar qualidade dos itens, distribuição de respostas, índice de acerto, itens com comportamento anômalo e aderência dos descritores à matriz curricular. Essa fase ainda pode ser considerada pré-TRI, pois busca formar o banco de dados necessário à calibração posterior.",
)
add_para(
    doc,
    "Na segunda etapa, com maior volume de respostas, os itens poderão ser analisados quanto à dificuldade e capacidade de discriminação entre estudantes com diferentes níveis de domínio. Itens muito fáceis, muito difíceis, ambíguos ou pouco relacionados ao descritor deverão ser revisados ou excluídos. A existência de itens âncora permitirá comparar aplicações distintas e acompanhar evolução entre entrada, percurso e saída.",
)
add_para(
    doc,
    "Na terceira etapa, a escala de proficiência técnica poderá ser consolidada por curso e, posteriormente, por eixo tecnológico. Essa decisão é metodologicamente importante: algumas competências são específicas de cada curso, enquanto outras podem ser comuns a um eixo, como leitura técnica, segurança, raciocínio lógico, interpretação de normas e resolução de problemas. O método deve preservar essas duas camadas: comparabilidade geral e especificidade profissional.",
)

heading(doc, "8.6 Educação híbrida, EaD e evidências práticas", 2)
add_para(
    doc,
    "A proposta distingue educação híbrida de Educação a Distância. A educação híbrida combina momentos presenciais, atividades mediadas por tecnologia, práticas supervisionadas e produção de evidências em diferentes espaços de aprendizagem. A EaD, por sua vez, pressupõe organização pedagógica com mediação remota mais estruturada, regras específicas de acompanhamento e critérios próprios de presencialidade quando exigidos. Para a EPT, essa distinção é decisiva, pois nem toda competência técnica pode ser validada apenas por atividade remota.",
)
add_para(
    doc,
    "No SATEC-CE/TRI, a escolha do instrumento dependerá da natureza da competência. Conteúdos conceituais podem ser avaliados por itens objetivos ou estudos de caso em ambiente digital. Procedimentos técnicos, uso de equipamentos, práticas de segurança, atendimento, montagem, manutenção, cuidado em saúde ou manejo produtivo devem gerar evidências presenciais, simuladas, laboratoriais, supervisionadas ou registradas em campo. Nas Escolas do Campo e na EFA, o tempo comunidade deve ser reconhecido como espaço legítimo de evidência, desde que acompanhado por protocolo, rubrica e validação docente.",
)
add_para(
    doc,
    "Essa arquitetura evita dois riscos. O primeiro é reduzir a EPT a uma prova teórica, incapaz de observar desempenho profissional. O segundo é tratar toda prática como evidência informal, sem critérios comuns de interpretação. A combinação entre itens, rubricas e portfólios permite avaliar conhecimento, execução e reflexão técnica.",
)

heading(doc, "8.7 Aspectos éticos e governança dos dados", 2)
add_para(
    doc,
    "A pesquisa lida com dados educacionais de estudantes e, por isso, deve adotar princípios de proteção, transparência e finalidade pedagógica. Os resultados individuais não devem ser usados para exposição pública, punição ou rotulação permanente. A finalidade principal é apoiar a aprendizagem e a melhoria institucional. Quando houver coleta para fins de pesquisa acadêmica, a participação deverá observar autorizações institucionais, consentimento informado e avaliação ética quando aplicável.",
)
add_para(
    doc,
    "A governança dos dados deverá prever quem coleta, quem acessa, por quanto tempo os dados são armazenados, quais informações são anonimizadas e como os relatórios são compartilhados. No uso de IA Generativa, recomenda-se evitar o envio de dados pessoais identificáveis a ferramentas externas. Quando a IA for usada para gerar atividades ou relatórios, deve receber dados agregados, descritores e níveis de proficiência, não informações sensíveis dos estudantes.",
)

heading(doc, "8.8 Riscos de implementação e medidas de mitigação", 2)
add_table(
    doc,
    ["Risco", "Efeito possível", "Mitigação"],
    [
        ["Transformar o método em exame classificatório", "Reduzir a função formativa e gerar pressão indevida.", "Definir finalidade diagnóstica, devolutiva pedagógica e plano de intervenção obrigatório."],
        ["Itens desalinhados ao currículo", "Medir conteúdos que não representam o curso.", "Validação por professores especialistas e revisão a partir das matrizes vigentes."],
        ["Baixa adesão docente", "Uso superficial dos relatórios.", "Formação continuada, coautoria dos instrumentos e relatórios claros."],
        ["Pouco volume de dados para TRI", "Impossibilidade de calibrar escala robusta no início.", "Iniciar com análise clássica, ampliar aplicações e criar banco progressivo."],
        ["Uso inadequado de IA", "Atividades genéricas, vieses ou exposição de dados.", "Revisão humana, anonimização e protocolo de uso responsável."],
    ],
)

heading(doc, "9 PRODUTO TÉCNICO-TECNOLÓGICO", 1)
add_para(
    doc,
    "O produto técnico-tecnológico resultante será o Caderno Metodológico SATEC-CE/TRI, acompanhado de instrumentos aplicáveis aos cursos-piloto. O produto deverá ser organizado para uso por professores, coordenadores e gestores, com linguagem técnica acessível e orientações práticas de implementação.",
)
add_table(
    doc,
    ["Componente do produto", "Descrição"],
    [
        ["Caderno metodológico", "Explica fundamentos, dimensões, etapas de aplicação, papéis da equipe e critérios de interpretação."],
        ["Matrizes avaliativas piloto", "Relacionam competências, descritores, componentes curriculares e evidências por curso."],
        ["Banco inicial de itens", "Reúne questões objetivas e situações-problema classificadas por descritor e nível de dificuldade esperado."],
        ["Rubricas práticas", "Avaliam execução técnica, segurança, comunicação, autonomia e resolução de problemas."],
        ["Protocolo de aplicação", "Define momentos de aplicação, devolutiva, recomposição e reavaliação."],
        ["Modelo de dashboard", "Apresenta indicadores por estudante, turma, curso, escola, eixo tecnológico e momento avaliativo."],
    ],
)
add_para(
    doc,
    "O produto final não deve ser entendido como prova única ou exame classificatório. Sua finalidade é formativa e diagnóstica: orientar a melhoria do ensino técnico, apoiar professores e estudantes e oferecer evidências para gestão pedagógica. A dimensão somativa pode existir na avaliação de saída, mas subordinada ao compromisso de aperfeiçoamento da formação profissional.",
)

heading(doc, "10 CRONOGRAMA", 1)
add_table(
    doc,
    ["Período", "Atividade principal", "Produto parcial"],
    [
        ["Mês 1", "Revisão bibliográfica e análise documental", "Quadro teórico e normativo"],
        ["Mês 2", "Seleção dos cursos-piloto e desenho da matriz avaliativa", "Matriz avaliativa preliminar"],
        ["Mês 3", "Elaboração de itens, situações-problema e rubricas", "Banco inicial de instrumentos"],
        ["Mês 4", "Validação por especialistas e ajustes", "Instrumentos revisados"],
        ["Mês 5", "Aplicação diagnóstica de entrada", "Relatório inicial"],
        ["Meses 6 e 7", "Intervenção pedagógica e trilhas de recomposição", "Planos por descritor"],
        ["Mês 8", "Avaliação de percurso", "Relatório comparativo"],
        ["Meses 9 e 10", "Ajuste metodológico e expansão controlada", "Versão 2 do método"],
        ["Mês 11", "Avaliação de saída", "Relatório final de proficiência"],
        ["Mês 12", "Redação final e consolidação do produto", "TCC e caderno metodológico"],
    ],
)

heading(doc, "11 RESULTADOS ESPERADOS", 1)
add_para(
    doc,
    "Espera-se que a pesquisa produza um método aplicável de avaliação técnica para cursos profissionalizantes da rede estadual do Ceará. Em termos pedagógicos, espera-se que o SATEC-CE/TRI permita identificar lacunas de aprendizagem de forma mais precisa, orientar intervenções por descritor, apoiar professores no planejamento e oferecer devolutivas compreensíveis aos estudantes.",
)
add_para(
    doc,
    "Em termos institucionais, espera-se que o método contribua para o acompanhamento da qualidade dos cursos técnicos, considerando não apenas desempenho discente, mas também condições de oferta, permanência, práticas profissionais e relação com o mundo do trabalho. Em termos científicos, espera-se contribuir para a discussão sobre uso da TRI e de Learning Analytics na Educação Profissional e Tecnológica, especialmente em redes públicas estaduais.",
)
add_para(
    doc,
    "Como resultado técnico, espera-se disponibilizar uma versão inicial do Caderno Metodológico SATEC-CE/TRI e seus instrumentos, possibilitando posterior ampliação para outros cursos, escolas e eixos tecnológicos. A médio prazo, o método poderá apoiar metas de melhoria da EPT no Ceará, induzindo decisões mais precisas sobre currículo, formação docente, infraestrutura, recomposição de aprendizagens e acompanhamento de egressos.",
)
add_table(
    doc,
    ["Indicador de sucesso", "Como será observado"],
    [
        ["Clareza dos descritores", "Percentual de professores que consideram os descritores compreensíveis e alinhados ao curso."],
        ["Utilidade pedagógica", "Uso dos relatórios para replanejamento, recomposição e agrupamento de estudantes."],
        ["Evolução dos estudantes", "Comparação entre avaliação de entrada, percurso e saída por descritor."],
        ["Qualidade dos itens", "Análise de dificuldade, discriminação e revisão por especialistas."],
        ["Escalabilidade", "Capacidade de adaptar o método para outros cursos e eixos tecnológicos."],
        ["Equidade", "Identificação de lacunas sem exposição indevida ou rotulação dos estudantes."],
    ],
)

heading(doc, "12 CONSIDERAÇÕES FINAIS", 1)
add_para(
    doc,
    "A proposta do SATEC-CE/TRI parte de uma convicção pedagógica: avaliar melhor é ensinar melhor. Quando a avaliação revela apenas notas, sua capacidade de transformar o ensino é limitada. Quando revela competências, descritores, níveis de proficiência e necessidades de intervenção, torna-se ferramenta de gestão pedagógica e justiça educacional.",
)
add_para(
    doc,
    "O Ceará possui uma rede de Educação Profissional com matrizes curriculares estruturadas, diversidade de cursos e experiências territoriais relevantes. O desafio colocado por esta pesquisa é transformar essa riqueza curricular em evidências sistemáticas de aprendizagem técnica. Para isso, a Teoria de Resposta ao Item, a avaliação por competências, Learning Analytics e Inteligência Artificial Generativa podem ser articuladas em um método aplicável, ético e orientado à melhoria.",
)
add_para(
    doc,
    "O TCC propõe, portanto, um caminho de pesquisa e intervenção em nível de mestrado profissional. O produto esperado não é apenas um documento, mas um instrumento de trabalho para professores, escolas e gestores: um método capaz de diagnosticar, acompanhar e aprimorar a formação técnica dos estudantes da rede estadual do Ceará.",
)

heading(doc, "REFERÊNCIAS", 1)
refs = [
    "BLOOM, B. S. The 2 sigma problem: the search for methods of group instruction as effective as one-to-one tutoring. Educational Researcher, v. 13, n. 6, p. 4-16, 1984.",
    "BRASIL. Lei nº 9.394, de 20 de dezembro de 1996. Estabelece as diretrizes e bases da educação nacional. Brasília: Presidência da República, 1996.",
    "BRASIL. Ministério da Educação. Base Nacional Comum Curricular. Brasília: MEC, 2018.",
    "BRASIL. Conselho Nacional de Educação. Resolução CNE/CP nº 1, de 5 de janeiro de 2021. Define as Diretrizes Curriculares Nacionais Gerais para a Educação Profissional e Tecnológica. Brasília: MEC/CNE, 2021.",
    "BRASIL. Lei nº 14.945, de 31 de julho de 2024. Altera a Lei nº 9.394/1996 e estabelece diretrizes para a política nacional de Ensino Médio. Brasília: Presidência da República, 2024.",
    "CEARÁ. Secretaria da Educação do Estado do Ceará. Caderno de Matrizes 2026: Escolas Estaduais de Educação Profissional. Fortaleza: SEDUC/COEDP/CEDET, 2026.",
    "CEARÁ. Secretaria da Educação do Estado do Ceará. Matrizes curriculares: Escolas de Ensino Médio do Campo e Escola Família Agrícola. Fortaleza: SEDUC-CE, 2025.",
    "CEARÁ. Secretaria da Educação do Estado do Ceará. Documento Curricular Referencial do Ceará. Fortaleza: SEDUC-CE, s.d.",
    "GAŠEVIĆ, D.; DAWSON, S.; SIEMENS, G. Let's not forget: learning analytics are about learning. TechTrends, v. 59, n. 1, p. 64-71, 2015.",
    "HOLMES, W.; BIALIK, M.; FADEL, C. Artificial Intelligence in Education: promises and implications for teaching and learning. Boston: Center for Curriculum Redesign, 2019.",
    "INEP. Avaliação da Educação Profissional e Tecnológica. Brasília: Instituto Nacional de Estudos e Pesquisas Educacionais Anísio Teixeira, 2026.",
    "LANG, C.; SIEMENS, G.; WISE, A.; GAŠEVIĆ, D. (org.). Handbook of Learning Analytics. New York: Society for Learning Analytics Research, 2017.",
    "LUCKIN, R.; HOLMES, W.; GRIFFITHS, M.; FORCIER, L. B. Intelligence Unleashed: an argument for AI in Education. London: Pearson, 2016.",
    "PASQUALI, L. Psicometria: teoria dos testes na psicologia e na educação. 5. ed. Petrópolis: Vozes, 2013.",
    "PASQUALI, L. TRI: Teoria de Resposta ao Item: teoria, procedimentos e aplicações. Curitiba: Appris, 2018.",
    "PASQUALI, L.; PRIMI, R. Fundamentos da teoria da resposta ao item: TRI. Avaliação Psicológica, v. 2, n. 2, p. 99-110, 2003.",
    "SIEMENS, G.; GAŠEVIĆ, D. Guest editorial: learning and knowledge analytics. Educational Technology & Society, v. 15, n. 3, p. 1-2, 2012.",
]
for ref in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(6)
    p.add_run(ref)

doc.add_page_break()
heading(doc, "APÊNDICE A - MATRIZ AVALIATIVA PILOTO: TÉCNICO EM INFORMÁTICA", 1)
add_para(
    doc,
    "A matriz a seguir exemplifica como um curso-piloto pode ser transformado em descritores avaliáveis. A versão final deverá ser validada por professores especialistas do curso e alinhada ao CNCT, ao DCRC e à matriz curricular vigente.",
)
add_table(
    doc,
    ["Componente", "Descritor avaliável", "Tipo de evidência", "Momento"],
    [
        ["Informática Básica", "Identificar componentes de hardware, periféricos, sistemas operacionais e práticas seguras de uso.", "Teste objetivo e tarefa prática", "Entrada"],
        ["Lógica de Programação", "Interpretar problemas, decompor etapas e representar algoritmos em pseudocódigo ou fluxograma.", "Situação-problema", "Percurso"],
        ["Arquitetura e Manutenção", "Diagnosticar falhas simples, propor procedimento e justificar cuidados de segurança.", "Rubrica prática", "Percurso"],
        ["HTML/CSS", "Construir página web simples com estrutura semântica e estilização coerente.", "Produto prático", "Percurso"],
        ["Banco de Dados", "Modelar entidades, atributos e relações para um problema contextualizado.", "Estudo de caso", "Percurso"],
        ["Programação Web", "Integrar lógica, interface e persistência em solução funcional simples.", "Projeto integrador", "Saída"],
        ["Redes de Computadores", "Interpretar topologia básica e propor configuração ou diagnóstico inicial.", "Situação-problema", "Saída"],
    ],
)

heading(doc, "APÊNDICE B - MODELO DE RELATÓRIO PEDAGÓGICO", 1)
add_para(
    doc,
    "O relatório pedagógico deverá apresentar informações em quatro níveis: estudante, turma, curso e escola. Para o estudante, recomenda-se indicar nível de proficiência, descritores consolidados, descritores em atenção e trilha de recomposição. Para a turma, recomenda-se apresentar distribuição por nível, descritores mais frágeis e sugestões de agrupamento. Para o curso, o relatório deve indicar pontos críticos por componente e evolução entre entrada, percurso e saída. Para a escola, deve apoiar planejamento, formação docente e gestão de infraestrutura.",
)
add_table(
    doc,
    ["Indicador", "Uso pedagógico"],
    [
        ["Proficiência estimada", "Interpretar posição do estudante na escala técnica."],
        ["Descritores críticos", "Planejar recomposição específica."],
        ["Evolução temporal", "Comparar entrada, percurso e saída."],
        ["Mapa da turma", "Formar grupos de intervenção."],
        ["Condições de oferta", "Relacionar desempenho a laboratório, insumos e carga prática."],
        ["Recomendação de IA", "Gerar atividades e trilhas, sempre revisadas pelo professor."],
    ],
)

heading(doc, "APÊNDICE C - ORIENTAÇÕES PARA USO RESPONSÁVEL DE IA", 1)
add_bullets(
    doc,
    [
        "A IA deve apoiar a produção de atividades, feedbacks e relatórios, mas não substituir a decisão docente.",
        "Dados pessoais dos estudantes devem ser minimizados, protegidos e tratados conforme normas aplicáveis.",
        "Toda recomendação gerada por IA deve passar por revisão humana.",
        "As trilhas de recomposição devem evitar rotular estudantes de forma fixa.",
        "A escola deve registrar critérios de uso, limites, responsáveis e formas de auditoria pedagógica.",
    ],
)

heading(doc, "APÊNDICE D - MODELO DE DESCRITOR E ITEM", 1)
add_para(
    doc,
    "O exemplo abaixo ilustra a lógica de construção de item para o curso Técnico em Informática. Trata-se de modelo preliminar, a ser validado por docentes do curso.",
)
add_table(
    doc,
    ["Campo", "Exemplo"],
    [
        ["Curso", "Técnico em Informática"],
        ["Componente", "Lógica de Programação"],
        ["Descritor", "Resolver problema simples por decomposição de etapas e representação algorítmica."],
        ["Nível esperado", "Nível 2 - Básico operacional"],
        ["Situação-problema", "Uma escola deseja calcular a média final de um estudante a partir de quatro notas e informar se ele foi aprovado, em recuperação ou reprovado."],
        ["Evidência", "Escolha da sequência lógica correta, uso de condição e interpretação do resultado."],
        ["Parâmetro inicial esperado", "Dificuldade média; item candidato a calibração após aplicação piloto."],
    ],
)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
