from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.text.paragraph import Paragraph
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
INPUT = ROOT / "norteadores" / "TCC_SIDEP-CE.docx"
OUTPUT = ROOT / "outputs" / "TCC_SIDEP-CE_mestrado_metodologia_v0_2.docx"


def insert_after(paragraph, text="", style=None):
    new_p = deepcopy(paragraph._p)
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    inserted.clear()
    inserted.style = style or "Normal"
    if text:
        inserted.add_run(text)
    return inserted


def clear_and_write(paragraph, text, style=None):
    paragraph.clear()
    if style:
        paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def add_para_after(anchor, text, style=None):
    p = insert_after(anchor, "", style)
    p.add_run(text)
    return p


def add_bullets_after(anchor, items):
    current = anchor
    for item in items:
        current = add_para_after(current, item, "List Bullet")
    return current


def find_para(doc, startswith):
    for p in doc.paragraphs:
        if p.text.strip().startswith(startswith):
            return p
    raise ValueError(f"Paragraph not found: {startswith}")


def find_heading(doc, startswith):
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading") and p.text.strip().startswith(startswith):
            return p
    raise ValueError(f"Heading not found: {startswith}")


def heading_level(paragraph):
    name = paragraph.style.name
    if not name.startswith("Heading"):
        return None
    try:
        return int(name.replace("Heading", "").strip())
    except ValueError:
        return None


def find_section_end(doc, heading_startswith):
    paragraphs = doc.paragraphs
    start = None
    level = None
    for i, p in enumerate(paragraphs):
        if p.style.name.startswith("Heading") and p.text.strip().startswith(heading_startswith):
            start = i
            level = heading_level(p)
            break
    if start is None:
        raise ValueError(f"Section not found: {heading_startswith}")

    end = start
    for j in range(start + 1, len(paragraphs)):
        current_level = heading_level(paragraphs[j])
        if current_level is not None and current_level <= level:
            break
        end = j
    return paragraphs[end]


def add_table_after(paragraph, rows, cols):
    tbl = paragraph._parent.add_table(rows=rows, cols=cols, width=Inches(6.5))
    paragraph._p.addnext(tbl._tbl)
    tbl.style = "Table Grid"
    return tbl


def apply_replacements(doc, replacements):
    def replace_in_paragraph(paragraph):
        for run in paragraph.runs:
            text = run.text
            for old, new in replacements.items():
                text = text.replace(old, new)
            run.text = text

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)


def main():
    doc = Document(INPUT)

    # Title update.
    title = find_para(doc, "AVALIA")
    title.clear()
    title.add_run(
        "SIDEP-CE: SISTEMA DE DIAGNOSTICO DA EDUCACAO PROFISSIONAL DO CEARA "
        "BASEADO EM COMPETENCIAS, TRI E INTELIGENCIA ARTIFICIAL GENERATIVA"
    ).bold = True

    # Strengthen chapter 7.
    cap7 = find_heading(doc, "7 PROPOSTA DO")
    p = add_para_after(
        cap7,
        "Nesta versao metodologica, o SIDEP-CE passa a ser tratado como matriz "
        "avaliativa evolutiva. A matriz inicial do curso Tecnico em Informatica, "
        "denominada Matriz SIDEP-CE v0.1, nao constitui um produto fechado, mas "
        "uma hipotese pedagogica fundamentada, a ser validada por aplicacao real, "
        "analise docente, evidencias de aprendizagem e estudos psicometricos progressivos.",
    )
    p = add_para_after(
        p,
        "A construcao da matriz deve seguir uma ordem metodologica obrigatoria: "
        "primeiro a leitura da matriz curricular e do perfil profissional do curso; "
        "depois a definicao das competencias amplas; em seguida o desdobramento em "
        "descritores observaveis; por fim, a vinculacao de itens, rubricas, projetos, "
        "relatorios e trilhas de recomposicao. Essa ordem evita que o banco de questoes "
        "determine artificialmente o curriculo avaliado.",
    )

    anchor7 = find_section_end(doc, "7.3 N")
    h = add_para_after(anchor7, "7.4 Matriz SIDEP-CE v0.1 para o curso Tecnico em Informatica", "Heading 2")
    p = add_para_after(
        h,
        "A matriz v0.1 do curso Tecnico em Informatica foi organizada com 10 competencias "
        "amplas, codificadas de C01 a C10, e 40 descritores avaliaveis, codificados de "
        "D01 a D40. As competencias representam grandes capacidades profissionais, "
        "enquanto os descritores traduzem essas capacidades em evidencias observaveis, "
        "capazes de orientar itens objetivos, situacoes-problema, tarefas praticas, "
        "rubricas de laboratorio, projetos integradores, estagio e atividades em tempo comunidade.",
    )
    p = add_para_after(
        p,
        "A matriz nao deve ser ampliada por intuicao isolada. Novas competencias ou "
        "descritores devem surgir apenas quando a aplicacao em campo, a analise dos "
        "resultados, a revisao de professores especialistas ou a matriz curricular "
        "indicarem lacunas reais. Assim, a versao v0.1 funciona como ponto de partida "
        "para validacao e nao como limite definitivo do metodo.",
    )

    h = add_para_after(p, "7.5 Competencia, descritor e indicador avaliativo", "Heading 2")
    p = add_para_after(
        h,
        "No SIDEP-CE, competencia, descritor e indicador possuem funcoes diferentes. "
        "Competencia e a capacidade ampla de mobilizar conhecimentos, habilidades, "
        "atitudes e procedimentos em situacoes tecnicas. Descritor e a evidencia "
        "observavel dessa competencia, escrita de forma especifica para orientar "
        "avaliacao e relatorio. Indicador ou subdescritor e uma camada interna de "
        "detalhamento, util ao planejamento docente e a recomposicao, sem necessidade "
        "de se tornar codigo principal da matriz estadual.",
    )
    p = add_para_after(
        p,
        "Desse modo, a questao mede prioritariamente um descritor; o conjunto de "
        "descritores interpreta uma competencia; e os indicadores apoiam a intervencao "
        "pedagogica fina. Essa arquitetura e coerente com a logica pre-TRI/TRI, pois "
        "permite acumular respostas por descritor antes de consolidar escalas mais "
        "amplas de proficiencia tecnica.",
    )

    h = add_para_after(p, "7.6 Regras para evolucao da matriz avaliativa", "Heading 2")
    p = add_para_after(
        h,
        "A criacao de novas competencias deve ocorrer quando a matriz curricular, o "
        "perfil profissional de conclusao, o CNCT ou a pratica docente indicarem uma "
        "capacidade profissional nao contemplada pelas competencias existentes. Ja a "
        "criacao de novos descritores deve ocorrer quando um descritor estiver amplo "
        "demais, quando itens vinculados medirem objetos diferentes, quando houver "
        "habilidade tecnica importante sem evidencia avaliavel ou quando os relatorios "
        "nao permitirem planejar recomposicao pedagogica.",
    )
    p = add_para_after(
        p,
        "Descritores tambem podem ser revisados, fundidos ou removidos quando se "
        "mostrarem repetitivos, pouco observaveis, excessivamente memoristicos, "
        "desalinhados da matriz ou com baixo valor pedagogico e psicometrico apos "
        "aplicacoes. Portanto, o SIDEP-CE trabalha com matriz viva, submetida a ciclos "
        "de estudo, aplicacao, analise e replanejamento.",
    )

    # Strengthen methodology.
    etapas = find_section_end(doc, "8.2 Etapas")
    h = add_para_after(etapas, "8.2.1 Protocolo de validacao da Matriz SIDEP-CE v0.1", "Heading 2")
    p = add_para_after(
        h,
        "A validacao da matriz do curso Tecnico em Informatica devera seguir um "
        "protocolo em ciclos. O primeiro ciclo corresponde a analise documental das "
        "matrizes curriculares, do percurso formativo, do DCRC, do CNCT e das normas "
        "da EPT. O segundo ciclo corresponde a revisao por professores especialistas, "
        "coordencao tecnica e gestao pedagogica. O terceiro ciclo corresponde a "
        "aplicacao diagnostica piloto, analise de respostas, revisao de itens e "
        "ajuste dos descritores. O quarto ciclo corresponde a reaplicacao formativa "
        "e consolidacao de evidencias para futura calibracao TRI.",
    )
    p = add_para_after(
        p,
        "Esse protocolo permite que a matriz seja defendida academicamente como "
        "instrumento em desenvolvimento validado por evidencias, e nao apenas como "
        "lista arbitraria de conteudos. O criterio de qualidade da matriz sera sua "
        "capacidade de gerar diagnosticos interpretaveis, relatorios acionaveis e "
        "intervencoes pedagogicas observaveis.",
    )

    analise = find_section_end(doc, "8.5 Estrat")
    h = add_para_after(analise, "8.5.1 Transicao da fase pre-TRI para a fase TRI", "Heading 2")
    p = add_para_after(
        h,
        "Na fase pre-TRI, o sistema utiliza dificuldade inicial estimada, percentual "
        "de acerto por descritor, consistencia pedagogica dos itens e analise de "
        "padroes de resposta. Essa fase e necessaria porque a TRI exige volume "
        "suficiente de respostas, estabilidade do banco, itens ancora e controle de "
        "qualidade dos instrumentos.",
    )
    p = add_para_after(
        p,
        "A transicao para TRI deve ocorrer somente quando houver dados suficientes "
        "para estimar parametros psicometricos com confiabilidade. Ate la, os relatorios "
        "devem indicar que os resultados usam parametros inspirados em TRI e analise "
        "classica de itens. Essa decisao fortalece a honestidade metodologica do projeto "
        "e evita apresentar como TRI uma escala ainda nao calibrada.",
    )

    gov = find_section_end(doc, "8.7 Aspectos")
    h = add_para_after(gov, "8.7.1 Regras de negocio e governanca pedagogica do sistema", "Heading 2")
    p = add_para_after(
        h,
        "O SIDEP-CE devera operar com separacao clara de perfis. O estudante acessa "
        "a avaliacao por codigo, informa nome completo ou autentica por cadastro "
        "institucional e responde em computador, tablet ou celular. Ele nao visualiza "
        "pesos, parametros pre-TRI/TRI, diagnostico bruto, analise da turma ou "
        "intervencoes automaticas.",
    )
    p = add_para_after(
        p,
        "O professor podera criar avaliacoes por uma ou mais disciplinas ou componentes, "
        "selecionando descritores e definindo a distribuicao de questoes. Cada avaliacao "
        "devera conter no minimo 20 questoes e no maximo 80, respeitando a disponibilidade "
        "de itens validados no banco. A gestao escolar visualizara relatorios agregados "
        "por turma, curso, escola e descritor, com finalidade de acompanhamento e apoio "
        "pedagogico, nao de exposicao ou punicao.",
    )
    p = add_para_after(
        p,
        "Toda questao devera estar vinculada a curso, componente curricular, competencia "
        "e descritor. O banco de itens devera registrar status do item, historico de "
        "aplicacao, desempenho, revisoes e possibilidade futura de calibracao. A IA "
        "Generativa podera apoiar a elaboracao de atividades, trilhas e relatorios, "
        "mas toda recomendacao devera passar por revisao humana.",
    )

    # Product chapter.
    produto = find_section_end(doc, "9 PRODUTO")
    p = add_para_after(
        produto,
        "Com a evolucao da proposta, o produto tecnico-tecnologico passa a ser "
        "compreendido em duas frentes integradas: o Caderno Metodologico SIDEP-CE e "
        "a plataforma digital SIDEP-CE. O caderno organiza a fundamentacao, matriz, "
        "regras de validacao, exemplos de itens e orientacoes de uso. A plataforma "
        "materializa esses principios em area do estudante, area do professor, area "
        "da gestao escolar, banco de itens, criador de avaliacoes, relatorios e trilhas "
        "de recomposicao.",
    )
    p = add_para_after(
        p,
        "No MVP atual, a base do curso Tecnico em Informatica conta com 10 competencias, "
        "40 descritores e 78 questoes piloto. Essa base deve ser tratada como matriz "
        "de validacao inicial, preparada para ampliacao responsavel a partir de novos "
        "dados, revisao docente e validacao psicometrica progressiva.",
    )

    # Append new appendices at the end of the document.
    doc.add_page_break()
    doc.add_heading(
        "APENDICE E - MATRIZ SIDEP-CE v0.1: COMPETENCIAS DO CURSO TECNICO EM INFORMATICA",
        level=1,
    )
    doc.add_paragraph(
        "A matriz v0.1 organiza o curso Tecnico em Informatica em 10 competencias amplas. "
        "Ela foi derivada da matriz curricular, do percurso formativo, do TCC SIDEP-CE "
        "e de referencia do CNCT para o eixo Informacao e Comunicacao."
    )

    competencias = [
        ("C01", "Operar recursos computacionais, sistemas operacionais e praticas de seguranca digital para uso responsavel da tecnologia."),
        ("C02", "Resolver problemas computacionais por meio de logica de programacao, algoritmos, programacao estruturada e orientacao a objetos."),
        ("C03", "Realizar diagnostico, montagem, configuracao, manutencao e suporte em computadores, perifericos e ambientes de hardware."),
        ("C04", "Desenvolver interfaces, paginas e aplicacoes web com estrutura semantica, estilo, interacao, usabilidade e integracao com dados."),
        ("C05", "Modelar, implementar e consultar bancos de dados, aplicando organizacao, persistencia, integridade e recuperacao de informacoes."),
        ("C06", "Compreender, configurar e diagnosticar redes de computadores, servicos de conectividade, enderecamento e compartilhamento de recursos."),
        ("C07", "Aplicar tecnologias emergentes, robotica e inteligencia artificial em solucoes tecnicas contextualizadas e eticamente orientadas."),
        ("C08", "Produzir artefatos digitais, interfaces e conteudos visuais com principios de design, composicao, acessibilidade e comunicacao."),
        ("C09", "Integrar conhecimentos tecnicos em laboratorios, projetos integradores e praticas profissionais, documentando solucoes e resultados."),
        ("C10", "Planejar carreira, estagio, organizacao do trabalho, empreendedorismo e inovacao em contextos produtivos e sociais da informatica."),
    ]
    tbl = doc.add_table(rows=len(competencias) + 1, cols=2)
    tbl.cell(0, 0).text = "Codigo"
    tbl.cell(0, 1).text = "Competencia"
    for r, (codigo, desc) in enumerate(competencias, start=1):
        tbl.cell(r, 0).text = codigo
        tbl.cell(r, 1).text = desc

    doc.add_heading("APENDICE F - DESCRITORES AVALIAVEIS DA MATRIZ SIDEP-CE v0.1", level=1)
    doc.add_paragraph(
        "Os 40 descritores a seguir constituem a primeira versao avaliavel da matriz. "
        "Cada descritor deve orientar itens, rubricas, relatorios e trilhas de recomposicao."
    )
    descritores = [
        ("D01", "C01", "Utilizar recursos basicos de computadores, aplicativos, internet e ferramentas de produtividade em situacoes escolares, profissionais e comunitarias."),
        ("D02", "C01", "Distinguir hardware, software, perifericos, arquivos, pastas e recursos de armazenamento local ou em nuvem."),
        ("D03", "C01", "Aplicar praticas de seguranca da informacao, cidadania digital, protecao de dados, senhas e uso responsavel da internet."),
        ("D04", "C01", "Reconhecer funcoes de sistemas operacionais, gerenciamento de arquivos, usuarios, permissoes, processos e recursos do sistema."),
        ("D05", "C02", "Decompor problemas em etapas, identificar entradas, processamentos e saidas e representar solucoes por algoritmos ou pseudocodigo."),
        ("D06", "C02", "Interpretar variaveis, tipos de dados, operadores, entrada, saida e conversoes em programas simples."),
        ("D07", "C02", "Analisar estruturas condicionais, estruturas de repeticao e fluxo de execucao em algoritmos e programas."),
        ("D08", "C02", "Aplicar listas, funcoes, modularizacao, validacao de dados e combinacao de estruturas na solucao de problemas."),
        ("D09", "C02", "Reconhecer classes, objetos, atributos, metodos, encapsulamento e relacoes basicas da programacao orientada a objetos."),
        ("D10", "C02", "Depurar erros, prever saidas, testar comportamentos e justificar decisoes em programas estruturados ou orientados a objetos."),
        ("D11", "C03", "Identificar componentes internos, perifericos, funcoes de hardware e relacoes entre pecas do computador."),
        ("D12", "C03", "Executar procedimentos de montagem, configuracao inicial, instalacao e verificacao de equipamentos computacionais."),
        ("D13", "C03", "Diagnosticar sintomas de falhas, lentidao, superaquecimento e propor manutencao preventiva ou corretiva."),
        ("D14", "C03", "Aplicar normas de seguranca, organizacao de bancada, registro tecnico e qualidade em atividades praticas de hardware."),
        ("D15", "C04", "Reconhecer estrutura semantica de paginas HTML, elementos, atributos, links, imagens e organizacao do conteudo."),
        ("D16", "C04", "Aplicar seletores, propriedades CSS, cores, responsividade, layout e estilizacao visual de interfaces web."),
        ("D17", "C04", "Implementar interacoes, validacoes e comportamentos basicos em paginas e aplicacoes web."),
        ("D18", "C04", "Organizar, publicar e manter conteudos em plataformas digitais, sites, blogs ou sistemas de gerenciamento de conteudo."),
        ("D19", "C04", "Integrar interface web, logica de aplicacao e persistencia de dados em solucoes simples voltadas a demandas reais."),
        ("D20", "C05", "Modelar entidades, atributos, relacionamentos e cardinalidades para representar problemas contextualizados."),
        ("D21", "C05", "Definir tabelas, campos, chaves primarias, chaves estrangeiras e regras basicas de integridade."),
        ("D22", "C05", "Interpretar e escrever comandos SQL basicos para inserir, consultar, atualizar e excluir dados."),
        ("D23", "C05", "Aplicar praticas de organizacao, backup, seguranca, consistencia e recuperacao de informacoes em bancos de dados."),
        ("D24", "C06", "Identificar topologias, equipamentos, meios de transmissao e funcoes basicas em redes de computadores."),
        ("D25", "C06", "Configurar enderecamento, conectividade, compartilhamento de recursos e servicos basicos de rede."),
        ("D26", "C06", "Diagnosticar problemas de conectividade, desempenho, acesso e seguranca em ambientes de rede."),
        ("D27", "C06", "Planejar solucoes simples de rede para escola, comunidade, associacao, cooperativa ou pequeno empreendimento."),
        ("D28", "C07", "Reconhecer sensores, atuadores, placas, comandos e principios de automacao aplicados a prototipos."),
        ("D29", "C07", "Aplicar logica de programacao em prototipos de robotica, automacao ou monitoramento de situacoes do territorio."),
        ("D30", "C07", "Compreender conceitos basicos de inteligencia artificial, dados, automacao, aprendizagem de maquina e IA generativa."),
        ("D31", "C07", "Avaliar usos eticos, limites, riscos e possibilidades da IA em projetos tecnicos, educacionais, produtivos e comunitarios."),
        ("D32", "C08", "Aplicar principios de composicao visual, cor, tipografia, alinhamento, contraste e hierarquia da informacao."),
        ("D33", "C08", "Selecionar formatos, ferramentas e procedimentos adequados para producao, edicao e exportacao de artefatos digitais."),
        ("D34", "C08", "Produzir materiais digitais acessiveis e comunicativos para escola, comunidade, movimentos, redes sociais ou projetos."),
        ("D35", "C09", "Levantar necessidades reais, definir requisitos e propor solucoes tecnologicas contextualizadas para escola ou comunidade."),
        ("D36", "C09", "Planejar, prototipar, testar, documentar e apresentar solucoes de software em projetos integradores."),
        ("D37", "C09", "Executar praticas laboratoriais de hardware ou infraestrutura com registro tecnico, seguranca, colaboracao e qualidade."),
        ("D38", "C10", "Relacionar organizacao do trabalho, tecnicas produtivas, colaboracao, responsabilidade e comunicacao profissional."),
        ("D39", "C10", "Planejar carreira, estagio, portfolio, postura profissional e estrategias de insercao no mundo do trabalho."),
        ("D40", "C10", "Elaborar propostas de inovacao, empreendedorismo, gestao de startup e solucoes digitais com impacto social ou produtivo."),
    ]
    tbl = doc.add_table(rows=len(descritores) + 1, cols=3)
    tbl.cell(0, 0).text = "Descritor"
    tbl.cell(0, 1).text = "Competencia"
    tbl.cell(0, 2).text = "Evidencia avaliavel"
    for r, row in enumerate(descritores, start=1):
        for c, value in enumerate(row):
            tbl.cell(r, c).text = value

    doc.add_heading("APENDICE G - REGRAS DE NEGOCIO DO SIDEP-CE", level=1)
    p = doc.add_paragraph(
        "As regras de negocio abaixo orientam a transformacao do metodo em plataforma "
        "digital aplicavel em escala escolar e estadual."
    )
    add_bullets_after(
        p,
        [
            "Toda questao deve estar vinculada a curso, componente curricular, competencia e descritor.",
            "Nenhuma questao deve ser aplicada sem status de validacao pedagogica.",
            "O estudante acessa a prova por codigo de avaliacao ou turma e nao visualiza pesos, parametros ou diagnostico bruto.",
            "O professor pode criar avaliacao com uma ou mais disciplinas/componentes, respeitando minimo de 20 e maximo de 80 questoes.",
            "Quando o banco nao possuir itens suficientes para um descritor, o sistema deve sinalizar necessidade de producao ou revisao de itens.",
            "Os relatorios devem existir em camadas: estudante, turma, curso, escola e gestao.",
            "A IA Generativa pode apoiar trilhas e relatorios, mas toda recomendacao deve ter revisao humana.",
            "A matriz deve ser revisada periodicamente com base em dados, parecer docente e resultados de aplicacao.",
        ],
    )

    for section in doc.sections:
        section.top_margin = section.top_margin

    for p in doc.paragraphs:
        for run in p.runs:
            if run.font.size is None:
                run.font.size = Pt(12)

    apply_replacements(
        doc,
        {
            "APENDICE": "APÊNDICE",
            "Tecnico": "Técnico",
            "TECNICO": "TÉCNICO",
            "Informatica": "Informática",
            "INFORMATICA": "INFORMÁTICA",
            "COMPETENCIAS": "COMPETÊNCIAS",
            "DESCRITORES AVALIAVEIS": "DESCRITORES AVALIÁVEIS",
            "NEGOCIO": "NEGÓCIO",
            "Competencia": "Competência",
            "competencia": "competência",
            "competencias": "competências",
            "Descritor": "Descritor",
            "descritor": "descritor",
            "descritores": "descritores",
            "avaliaveis": "avaliáveis",
            "avaliavel": "avaliável",
            "validacao": "validação",
            "evolucao": "evolução",
            "hipotese": "hipótese",
            "pedagogica": "pedagógica",
            "questoes": "questões",
            "questao": "questão",
            "minimo": "mínimo",
            "maximo": "máximo",
            "transicao": "transição",
            "Transicao": "Transição",
            "Codigo": "Código",
            "Evidencia": "Evidência",
            "educacao": "educação",
            "negocio": "negócio",
            "governanca": "governança",
            "revisao": "revisão",
            "docente": "docente",
            "aplicacao": "aplicação",
            "curricular": "curricular",
            "analise": "análise",
            "profissional": "profissional",
            "estagio": "estágio",
            "robotica": "robótica",
            "inteligencia": "inteligência",
            "generativa": "generativa",
            "sistema": "sistema",
            "proficiencia": "proficiência",
        },
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
