from pathlib import Path

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
INPUT = ROOT / "outputs" / "TCC_SIDEP-CE_mestrado_metodologia_v0_2.docx"
OUTPUT = ROOT / "outputs" / "TCC_SIDEP-CE_mestrado_metodologia_v0_3.docx"


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def main():
    doc = Document(INPUT)

    doc.add_page_break()
    doc.add_heading("APENDICE H - ATUALIZACAO METODOLOGICA V0.3: CURADORIA DOCENTE E AREA ESCOLA", level=1)

    doc.add_paragraph(
        "A versao v0.3 do SIDEP-CE incorpora ao produto tecnico-tecnologico uma camada "
        "mais robusta de governanca do banco de itens e define a proxima etapa de "
        "desenvolvimento da area Escola/Gestao Escolar. Essa atualizacao fortalece a "
        "coerencia entre metodologia academica, regras de negocio e implementacao do MVP."
    )

    doc.add_heading("H.1 Situacao atual do MVP", level=2)
    doc.add_paragraph(
        "O nucleo da area do professor foi consolidado no MVP React. A plataforma passou "
        "a organizar competencias, descritores, questoes, curadoria docente e criacao de "
        "avaliacoes em um fluxo unico. O banco piloto do curso Tecnico em Informatica "
        "conta com 10 competencias, 40 descritores e 441 questoes, sendo 78 inicialmente "
        "validadas e 363 em revisao docente."
    )
    add_bullets(
        doc,
        [
            "competencias codificadas no padrao C01, C02, C03 e assim sucessivamente;",
            "descritores codificados no padrao D01, D02, D03 e assim sucessivamente;",
            "questoes vinculadas obrigatoriamente a descritores;",
            "recuperacao da competencia a partir do descritor vinculado;",
            "cadastro de alternativas, gabarito, justificativa pedagogica e dificuldade inicial pre-TRI;",
            "criador de avaliacoes usando exclusivamente questoes validadas;",
            "separacao entre area do estudante, area do professor e area de gestao.",
        ],
    )

    doc.add_heading("H.2 Status pedagogico das questoes", level=2)
    doc.add_paragraph(
        "A plataforma passa a reconhecer tres status principais de curadoria do item. "
        "Essa classificacao e essencial para impedir que questoes ainda nao revisadas "
        "cheguem ao estudante e para organizar o trabalho docente de validacao."
    )
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Status"
    table.cell(0, 1).text = "Sentido pedagogico"
    table.cell(0, 2).text = "Pode compor avaliacao?"
    rows = [
        ("rascunho", "Item em elaboracao ou devolvido para ajuste.", "Nao"),
        ("em_revisao", "Item aguardando curadoria docente.", "Nao"),
        ("validada", "Item revisado e liberado para uso pedagogico.", "Sim"),
    ]
    for index, row in enumerate(rows, start=1):
        for col, value in enumerate(row):
            table.cell(index, col).text = value

    doc.add_paragraph(
        "A regra de negocio adotada e que somente questoes com status validada podem "
        "ser usadas pelo criador de avaliacoes. Itens em revisao ou rascunho permanecem "
        "registrados no banco, mas ficam bloqueados para aplicacao."
    )

    doc.add_heading("H.3 Fluxo de validacao docente", level=2)
    doc.add_paragraph(
        "O fluxo de validacao docente foi incorporado ao MVP para transformar o banco "
        "de itens em um espaco de curadoria pedagogica, e nao apenas em um deposito de "
        "questoes."
    )
    add_numbered(
        doc,
        [
            "A questao e cadastrada ou importada no Banco de Itens.",
            "O item permanece como rascunho ou em revisao.",
            "O professor ou revisor abre a fila de validacao docente.",
            "O sistema permite visualizar a questao completa em uma janela de leitura.",
            "A janela exibe enunciado, alternativas, gabarito destacado, justificativa, competencia, descritor, componente e dificuldade pre-TRI.",
            "O revisor decide entre validar, devolver para revisao ou marcar como rascunho.",
            "Somente apos validacao o item pode compor avaliacoes dos estudantes.",
        ],
    )

    doc.add_heading("H.4 Cobertura do banco por competencia e descritor", level=2)
    doc.add_paragraph(
        "A versao v0.3 tambem introduz uma leitura de cobertura do banco de itens por "
        "competencia e por descritor. Essa leitura mostra totais de questoes, itens "
        "validados, itens em revisao e rascunhos."
    )
    add_bullets(
        doc,
        [
            "identificar competencias com baixa cobertura avaliativa;",
            "localizar descritores que ainda nao possuem quantidade suficiente de itens validados;",
            "planejar producao e revisao de novas questoes;",
            "preparar a futura analise classica dos itens;",
            "organizar o caminho para calibracao TRI apos volume suficiente de respostas.",
        ],
    )

    doc.add_heading("H.5 Relacao com a fase pre-TRI e TRI", level=2)
    doc.add_paragraph(
        "A existencia de uma fila de curadoria e de indicadores de cobertura nao configura "
        "ainda TRI plena. Esses mecanismos fazem parte da fase pre-TRI, pois organizam "
        "o banco, melhoram a qualidade pedagogica dos itens e preparam o historico de "
        "aplicacoes. A passagem para TRI exigira respostas reais, estabilidade do banco, "
        "itens-ancora, analise de dificuldade, discriminacao e validacao psicometrica."
    )

    doc.add_heading("H.6 Proxima etapa: area Escola/Gestao Escolar", level=2)
    doc.add_paragraph(
        "Com a area do professor consolidada como nucleo do MVP, a proxima sprint deve "
        "concentrar-se na area Escola/Gestao Escolar. Essa area devera permitir que a "
        "gestao acompanhe a aplicacao do SIDEP-CE em nivel institucional, sem invadir "
        "o espaco de autoria e curadoria do professor."
    )
    add_bullets(
        doc,
        [
            "visualizar dados da escola identificada pelo codigo INEP;",
            "acompanhar turmas, cursos, professores e avaliacoes vinculadas a escola;",
            "monitorar avaliacoes em rascunho, publicadas, em aplicacao e encerradas;",
            "acompanhar participacao dos estudantes por turma e curso;",
            "visualizar resultados agregados por turma, curso, competencia e descritor;",
            "identificar descritores criticos recorrentes;",
            "registrar ou acompanhar intervencoes pedagogicas planejadas;",
            "respeitar permissoes, LGPD e separacao entre relatorio pedagogico e exposicao indevida de estudantes.",
        ],
    )

    doc.add_heading("H.7 Regras iniciais da area Escola", level=2)
    add_bullets(
        doc,
        [
            "a escola deve possuir codigo INEP unico;",
            "a escola deve estar vinculada a CREDE ou SEFOR;",
            "a escola deve possuir tipo institucional parametrizavel, como EEEP, EEMCP, EEMTI, EEM, EJA, Escola do Campo ou EFA;",
            "a gestao escolar deve visualizar apenas dados da propria escola;",
            "perfis CREDE/SEFOR e SEDUC devem visualizar dados agregados conforme permissao;",
            "relatorios nominais devem ser restritos a perfis autorizados;",
            "a gestao escolar nao deve alterar gabaritos, respostas ou parametros de item sem permissao especifica;",
            "toda acao sensivel deve ser preparada para log de auditoria na arquitetura com banco real.",
        ],
    )

    doc.add_paragraph(
        "Essa etapa aproxima o SIDEP-CE de uma plataforma de gestao pedagogica estadual, "
        "pois desloca o olhar da prova isolada para a escola como unidade de acompanhamento, "
        "planejamento, recomposicao e melhoria continua da Educacao Profissional."
    )

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.size is None:
                run.font.size = Pt(12)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
