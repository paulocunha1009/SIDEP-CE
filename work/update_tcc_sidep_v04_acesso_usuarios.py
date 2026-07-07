from pathlib import Path

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
INPUT = ROOT / "outputs" / "TCC_SIDEP-CE_mestrado_metodologia_v0_3.docx"
OUTPUT = ROOT / "outputs" / "TCC_SIDEP-CE_mestrado_metodologia_v0_4.docx"


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main():
    doc = Document(INPUT)

    doc.add_page_break()
    doc.add_heading("APENDICE I - CONTROLE DE ACESSO, PERFIS E GOVERNANCA DE USUARIOS", level=1)

    doc.add_paragraph(
        "A versao v0.4 do SIDEP-CE acrescenta ao produto tecnico-tecnologico uma camada "
        "inicial de controle de acesso por usuario, senha e escopo institucional. Essa "
        "camada e indispensavel para que o sistema deixe de funcionar como prototipo aberto "
        "e passe a representar uma plataforma educacional com separacao de responsabilidades, "
        "seguranca e governanca de dados."
    )

    doc.add_heading("I.1 Principio geral", level=2)
    doc.add_paragraph(
        "O controle de acesso do SIDEP-CE deve seguir o principio do menor acesso necessario. "
        "Cada usuario deve visualizar apenas as informacoes relacionadas a sua funcao, sua escola, "
        "sua regional ou sua responsabilidade institucional. Essa diretriz reduz exposicao indevida "
        "de dados e esta alinhada a uma perspectiva pedagogica e etica de uso dos resultados."
    )

    doc.add_heading("I.2 Fluxo de acesso do estudante", level=2)
    doc.add_paragraph(
        "O estudante nao deve possuir acesso administrativo ao sistema. Na tela principal, o acesso "
        "do aluno foi colocado em destaque, com dois campos centrais: codigo da prova e nome completo. "
        "Essa decisao simplifica a experiencia do estudante e evita exposicao de diagnosticos, pesos, "
        "parametros pre-TRI/TRI, banco de itens ou relatorios."
    )

    doc.add_heading("I.3 Credenciais iniciais por perfil", level=2)
    table = doc.add_table(rows=6, cols=4)
    table.style = "Table Grid"
    headers = ["Perfil", "Usuario", "Senha inicial", "Observacao"]
    for col, value in enumerate(headers):
        table.cell(0, col).text = value
    rows = [
        ("Aluno", "Nao se aplica", "Nao se aplica", "Acessa por codigo da prova e nome completo."),
        ("Gestao escolar", "Codigo INEP", "Codigo INEP", "Troca obrigatoria no primeiro acesso."),
        ("Professor tecnico", "E-mail institucional", "CPF cadastrado", "Troca obrigatoria no primeiro acesso."),
        ("CREDE/SEFOR", "Codigo regional", "Senha institucional", "Acesso ao recorte da regional."),
        ("Administrador", "Usuario administrativo", "Senha administrativa", "Acesso total e redefinicao de senhas."),
    ]
    for row_index, row in enumerate(rows, start=1):
        for col, value in enumerate(row):
            table.cell(row_index, col).text = value

    doc.add_heading("I.4 Escopo de permissao", level=2)
    add_bullets(
        doc,
        [
            "Aluno: acessa apenas a avaliacao pelo codigo da prova e pelo nome completo.",
            "Professor: acessa a area professor, banco de itens, avaliacoes e relatorios do proprio escopo.",
            "Gestao escolar: acessa professores, avaliacoes e relatorios da propria escola.",
            "CREDE/SEFOR: acessa escolas, professores, banco de itens, avaliacoes e relatorios da propria regional.",
            "SEDUC: acessa indicadores agregados da rede.",
            "Administrador: acessa todas as abas e e o unico perfil autorizado a redefinir senhas.",
        ],
    )

    doc.add_heading("I.5 Redefinicao de senha", level=2)
    doc.add_paragraph(
        "A redefinicao de senha deve ser uma acao restrita ao Administrador. Quando a senha da "
        "gestao escolar e redefinida, ela retorna ao codigo INEP da escola. Quando a senha do "
        "professor e redefinida, ela retorna ao CPF cadastrado. Em ambos os casos, o usuario deve "
        "alterar a senha no proximo acesso."
    )

    doc.add_heading("I.6 Implicacoes para a arquitetura estadual", level=2)
    doc.add_paragraph(
        "No MVP local, as senhas sao utilizadas apenas para validar o fluxo de negocio. Em uma "
        "implantacao real, esse modelo deve ser migrado para autenticacao institucional, Supabase Auth "
        "ou outro provedor seguro, com armazenamento de senha por hash, politicas de recuperacao, "
        "logs de auditoria, bloqueio de contas inativas e permissoes em nivel de linha no banco de dados."
    )
    add_bullets(
        doc,
        [
            "A tabela de escola deve registrar politica de primeiro acesso da gestao escolar.",
            "A tabela de professor deve registrar CPF, perfil de acesso e politica de primeiro acesso.",
            "Relatorios devem respeitar filtros por escola, CREDE/SEFOR e professor responsavel.",
            "Acoes sensiveis, como redefinicao de senha e desativacao de usuario, devem gerar log de auditoria.",
            "Dados de estudantes devem permanecer protegidos e visiveis apenas para perfis autorizados.",
        ],
    )

    doc.add_paragraph(
        "Com essa camada, o SIDEP-CE passa a se aproximar de um sistema profissional de avaliacao "
        "educacional, pois combina matriz avaliativa, banco de itens, curadoria docente e governanca "
        "de acesso por responsabilidade institucional."
    )

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.size is None:
                run.font.size = Pt(12)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
