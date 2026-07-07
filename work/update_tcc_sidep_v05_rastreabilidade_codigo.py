from pathlib import Path

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
INPUT = ROOT / "outputs" / "TCC_SIDEP-CE_mestrado_metodologia_v0_4.docx"
OUTPUT = ROOT / "outputs" / "TCC_SIDEP-CE_mestrado_metodologia_v0_5.docx"


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main():
    doc = Document(INPUT)

    doc.add_page_break()
    doc.add_heading("APENDICE J - RASTREABILIDADE DA APLICACAO E CODIGO AVALIATIVO", level=1)

    doc.add_paragraph(
        "A versao v0.5 do SIDEP-CE acrescenta uma regra de governanca da aplicacao avaliativa: "
        "o codigo de acesso da avaliacao passa a ser gerado automaticamente pelo sistema e, apos "
        "a abertura da avaliacao, torna-se chave imutavel de rastreabilidade. Essa decisao nao e "
        "apenas operacional; ela tem relevancia metodologica para a pesquisa, pois protege a "
        "integridade da base de respostas e dos relatorios por turma, descritor, competencia, "
        "componente curricular e escola."
    )

    doc.add_heading("J.1 Fundamentacao da regra", level=2)
    doc.add_paragraph(
        "Em um sistema diagnostico baseado em competencias, descritores, banco de itens e futura "
        "calibracao psicometrica, cada aplicacao precisa ser identificavel de forma unica. O codigo "
        "da avaliacao conecta o estudante a uma prova especifica, mas tambem conecta a prova ao "
        "professor responsavel, a turma, a escola, a etapa avaliativa, ao conjunto de itens, as "
        "respostas enviadas e aos relatorios gerados."
    )
    doc.add_paragraph(
        "Se um codigo pudesse ser editado ou reaproveitado, haveria risco de mistura de respostas "
        "entre aplicacoes distintas, fragilizando a confiabilidade dos indicadores pedagogicos e "
        "das analises futuras. Por isso, o codigo deve ser tratado como identificador de aplicacao "
        "e nao como simples senha de acesso."
    )

    doc.add_heading("J.2 Regras de negocio incorporadas", level=2)
    add_bullets(
        doc,
        [
            "O codigo da avaliacao e gerado automaticamente pelo sistema, com token randomico.",
            "O professor pode gerar um novo codigo antes de abrir a avaliacao.",
            "O professor nao deve digitar ou editar manualmente o codigo.",
            "Ao abrir a avaliacao, o codigo fica bloqueado e imutavel.",
            "O codigo passa a vincular respostas, turma, professor, escola, etapa, status e relatorios.",
            "Avaliacao criada por engano pode ser excluida somente se ainda nao houver respostas registradas.",
            "Avaliacao com respostas nao pode ser excluida; deve ser encerrada ou corrigida.",
            "O codigo de avaliacao excluida permanece bloqueado e nao pode ser reaproveitado.",
        ],
    )

    doc.add_heading("J.3 Implicacoes para relatorios e auditoria", level=2)
    doc.add_paragraph(
        "A imutabilidade do codigo favorece a auditoria da aplicacao, pois permite reconstruir "
        "o percurso de uma avaliacao desde sua criacao ate a geracao dos relatorios. Essa regra "
        "tambem apoia a leitura longitudinal, pois evita que um mesmo identificador seja usado "
        "em momentos ou turmas diferentes."
    )
    add_bullets(
        doc,
        [
            "Relatorios individuais podem ser vinculados com seguranca a uma aplicacao especifica.",
            "Relatorios por turma podem distinguir diagnostico, formativo e final sem colisao de dados.",
            "Analises por descritor e competencia preservam o conjunto de itens efetivamente aplicado.",
            "A gestao escolar e regional pode auditar status, exclusoes e encerramentos.",
            "A base de respostas fica mais consistente para analise classica e futura TRI.",
        ],
    )

    doc.add_heading("J.4 Relevancia para a pesquisa de mestrado", level=2)
    doc.add_paragraph(
        "Para o produto tecnico-tecnologico do mestrado, essa regra fortalece a validade operacional "
        "do SIDEP-CE. A pesquisa nao depende apenas da existencia de questoes ou relatorios, mas da "
        "garantia de que cada evidencia de aprendizagem pertence a uma aplicacao unica e rastreavel. "
        "Assim, o codigo imutavel contribui para a confiabilidade dos dados utilizados na tomada de "
        "decisao pedagogica e na avaliacao da propria metodologia."
    )

    doc.add_heading("J.5 Recomendacao para banco real", level=2)
    doc.add_paragraph(
        "Na migracao para Supabase/PostgreSQL, recomenda-se criar uma tabela especifica para codigos "
        "bloqueados ou historico de aplicacoes, registrando codigo, avaliacao original, motivo do "
        "bloqueio, usuario responsavel, data e metadados de auditoria. Essa estrutura evita "
        "reaproveitamento indevido e prepara o sistema para uso institucional em escala estadual."
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Elemento", "Regra", "Finalidade"]
    for col, value in enumerate(headers):
        table.cell(0, col).text = value
    rows = [
        ("Codigo gerado", "Automatico e randomico", "Evitar duplicidade e reduzir erro humano"),
        ("Codigo aberto", "Bloqueado e imutavel", "Preservar rastreabilidade da aplicacao"),
        ("Avaliacao sem respostas", "Pode ser excluida", "Corrigir erro operacional sem contaminar dados"),
        ("Codigo excluido", "Nao pode ser reaproveitado", "Evitar colisao historica e auditoria ambigua"),
        ("Avaliacao com respostas", "Nao pode ser excluida", "Preservar relatorios e base empirica"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for col, value in enumerate(row):
            cells[col].text = value

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.size is None:
                run.font.size = Pt(12)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
