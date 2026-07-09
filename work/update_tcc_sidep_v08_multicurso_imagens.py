from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "outputs" / "TCC_SIDEP-CE_mestrado_metodologia_v0_7.docx"
OUT = ROOT / "outputs" / "TCC_SIDEP-CE_mestrado_metodologia_v0_8.docx"


def set_run(run, bold=False, color=None, size=None):
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    if size:
        run.font.size = Pt(size)


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        set_run(run, bold=True, color=(0, 48, 104), size=16 if level == 1 else 13)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def main():
    doc = Document(SRC)
    doc.add_page_break()

    add_heading(doc, "Adendo metodologico v0.8 - Governanca multcurso, avaliacoes e imagem opcional", 1)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Atualizacao incorporada ao SIDEP-CE em 09/07/2026. ")
    set_run(run, bold=True, color=(0, 120, 68))
    paragraph.add_run(
        "Este adendo registra aprimoramentos relevantes para a pesquisa de mestrado, mantendo foco na validade pedagogica, "
        "na rastreabilidade dos dados e na expansao estadual do metodo."
    )

    add_heading(doc, "1. Governanca multcurso", 2)
    doc.add_paragraph(
        "O SIDEP-CE passou a operar com selecao explicita de curso tecnico nas subtelas de competencias, descritores e questoes. "
        "Essa decisao prepara o sistema para receber novos cursos da Educacao Profissional sem provocar conflito entre codigos pedagogicos semelhantes."
    )
    add_bullets(
        doc,
        [
            "Cada curso pode possuir suas proprias competencias C01, C02, C03 e seus proprios descritores D01, D02, D03.",
            "Na interface, o professor continua lendo codigos simples e pedagogicos.",
            "Internamente, o sistema escopa codigos por curso quando necessario para preservar integridade do banco.",
            "O cadastro de questoes segue a sequencia propria do curso, evitando sobrescrita ou colisao de chaves.",
        ],
    )

    add_heading(doc, "2. Fluxo curso, componente, descritor e questao", 2)
    doc.add_paragraph(
        "O cadastro de questoes foi reorganizado para obedecer a cadeia pedagogica correta: curso tecnico, componente curricular, descritor vinculado e questao. "
        "Ao selecionar um componente, o sistema apresenta apenas os descritores daquele componente, evitando que o item seja associado a matriz inadequada."
    )
    add_bullets(
        doc,
        [
            "O codigo da questao e gerado automaticamente a partir da ultima questao cadastrada no curso.",
            "A questao so pode ser salva quando o descritor pertence ao componente selecionado.",
            "O sistema preserva a exigencia de status validada para uso em avaliacao.",
            "O banco continua bloqueando questoes duplicadas ou com enunciado/contexto equivalente.",
        ],
    )

    add_heading(doc, "3. Governanca das avaliacoes por perfil", 2)
    doc.add_paragraph(
        "A aba Avaliacoes foi ajustada para respeitar o escopo institucional de acesso, sem alterar a logica de selecao de componentes, descritores e quantidade de questoes."
    )
    add_bullets(
        doc,
        [
            "Professor visualiza e cria avaliacoes apenas nos cursos vinculados ao seu cadastro.",
            "Gestao escolar visualiza as avaliacoes da propria escola.",
            "CREDE/SEFOR visualiza as avaliacoes das escolas vinculadas a sua regional.",
            "SEDUC e Administrador visualizam todas as avaliacoes da rede.",
            "A listagem de aplicacoes exibe codigo, curso, turma, componentes, escola/INEP, professor/matricula e status.",
        ],
    )

    add_heading(doc, "4. Imagem opcional em itens", 2)
    doc.add_paragraph(
        "O sistema passou a admitir imagem opcional no cadastro de questoes. O recurso amplia a capacidade diagnostica em cursos tecnicos, pois permite itens baseados em diagramas, telas, tabelas, esquemas, mapas, fluxogramas e situacoes-problema visuais."
    )
    add_bullets(
        doc,
        [
            "A imagem nao e obrigatoria: questoes textuais continuam funcionando normalmente.",
            "Quando cadastrada, a imagem aparece na prova do estudante e no modal de validacao docente.",
            "A interface foi preparada para exibir imagens de forma responsiva em computador, tablet e celular.",
            "No piloto local, a imagem pode ser informada por URL publica ou base64.",
        ],
    )

    add_heading(doc, "5. Estrategia recomendada para Supabase gratuito", 2)
    doc.add_paragraph(
        "Para o uso online, a estrategia tecnicamente mais adequada e armazenar arquivos de imagem no Supabase Storage e gravar no PostgreSQL apenas a URL ou o caminho do arquivo. "
        "Essa decisao preserva a cota do banco, facilita backup e melhora a escalabilidade do piloto."
    )
    add_bullets(
        doc,
        [
            "Base64 funciona em testes locais, mas aumenta muito o tamanho do registro no banco.",
            "Supabase Storage e mais apropriado para imagens do banco de itens.",
            "No plano gratuito, o uso deve ser controlado por tamanho de arquivo, compressao e limpeza periodica.",
            "A implementacao online definitiva deve criar bucket especifico para imagens, politicas de leitura e campo persistente de URL da imagem na tabela de questoes.",
        ],
    )

    add_heading(doc, "6. Relevancia para a pesquisa de mestrado", 2)
    doc.add_paragraph(
        "Os aprimoramentos reforcam a tese de que o SIDEP-CE nao e apenas um aplicador de provas, mas um metodo de diagnostico tecnico baseado em matriz curricular, descritores observaveis, governanca de itens, rastreabilidade de aplicacoes e leitura pedagogica dos resultados."
    )
    add_bullets(
        doc,
        [
            "A regra multcurso permite expansao estadual sem perda de consistencia.",
            "A selecao por componente e descritor aumenta validade de conteudo dos itens.",
            "A hierarquia de acesso protege dados pedagogicos e organiza responsabilidade institucional.",
            "O suporte a imagem aproxima a avaliacao das praticas reais da Educacao Profissional.",
            "A persistencia online cria condicoes para coleta empirica, analise classica e futura calibracao TRI.",
        ],
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
